# EPU Builder v1.3.2 – Streamlit + SQLite
# Correzioni incluse:
# - RIMOSSI i blocchi duplicati "CREA NUOVA VOCE" e "ELENCO + DETTAGLIO VOCE" fuori dalle funzioni (causavano NameError su 'cap')
# - PRAGMA foreign_keys=ON su ogni connessione
# - Indici SQLite per performance
# - ui_clienti(): CAP rinominato in cap_zip per evitare ombreggiamento con 'cap' (capitoli)
# - import_materiali_csv(): parsing float robusto (virgole decimali)
# - export_preventivo_docx(): gestione sicura campo 'note'
# - Chiavi Streamlit già coerenti in ui_voci()
# - NEW: CSS per rimuovere “fullscreen” su tabelle, utility testo/filtri, blocco duplicati fornitori,
#        messaggio capitolo defaults aggiornato

import io
import re
import sqlite3
from contextlib import contextmanager
from typing import Optional, Dict

import pandas as pd
import streamlit as st

DB_PATH = "epu.db"
UM_CHOICES = ["Mt", "Mtq2", "Hr", "Nr", "Lt", "GG", "KG", "QL", "AC"]

st.set_page_config(page_title="EPU Builder v1.3.2", layout="wide")

# ------------------------------------------------------------------
# CSS globale: rimuovi pulsante "View fullscreen" sui dataframe/editor
# ------------------------------------------------------------------
_HF_CSS = """
<style>
button[kind="header"] svg[aria-label="View fullscreen"] { display: none !important; }
div[data-testid="stElementToolbar"] button[title="View fullscreen"] { display: none !important; }
/* Per nascondere tutta la toolbar degli elementi Streamlit, scommenta: */
/* div[data-testid="stElementToolbar"] { display: none !important; } */
</style>
"""
def inject_global_css():
    st.markdown(_HF_CSS, unsafe_allow_html=True)

inject_global_css()

# ------------------------------------------------------------------
# Utils
# ------------------------------------------------------------------
def _exec(con, sql, params=None):
    cur = con.cursor()
    cur.execute(sql, params or [])
    return cur

def _to_float(x, default=0.0):
    """Cast robusto con supporto alla virgola decimale."""
    if pd.isna(x):
        return default
    try:
        return float(str(x).replace(",", "."))
    except Exception:
        return default

# --- Utility testo/filtri (serviranno anche per i filtri stile Excel) ---
def _norm_text(x: str) -> str:
    """minuscolo, spazi singoli, rimuove punteggiatura semplice: utile per confronti su nomi."""
    if pd.isna(x):
        return ""
    x = re.sub(r"[^\w\s]", "", str(x).strip(), flags=re.UNICODE)
    x = re.sub(r"\s+", " ", x).strip().lower()
    return x

def _digits_only(x: str) -> str:
    """solo cifre (es. per P.IVA)."""
    return re.sub(r"\D", "", str(x or ""))

def like_mask(series: pd.Series, needle: str) -> pd.Series:
    """Filtro 'contains' case-insensitive; True se needle è vuoto."""
    if not needle:
        return pd.Series([True]*len(series))
    return series.fillna("").astype(str).str.contains(str(needle), case=False, regex=False)

# ------------------------------------------------------------------
# DB init
# ------------------------------------------------------------------
def init_db():
    with sqlite3.connect(DB_PATH) as con:
        cur = con.cursor()

        # Tabelle di dominio
        cur.execute("""
        CREATE TABLE IF NOT EXISTS categorie (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL UNIQUE
        )""")
        cur.execute("""
        CREATE TABLE IF NOT EXISTS fornitori (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL UNIQUE,
            piva TEXT,
            indirizzo TEXT,
            email TEXT,
            telefono TEXT
        )""")

        # Materiali
        cur.execute("""
        CREATE TABLE IF NOT EXISTS materiali_base (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            categoria_id INTEGER NOT NULL,
            fornitore_id INTEGER NOT NULL,
            codice_fornitore TEXT NOT NULL,
            descrizione TEXT NOT NULL,
            unita_misura TEXT NOT NULL,
            quantita_default REAL DEFAULT 1.0,
            prezzo_unitario REAL NOT NULL,
            FOREIGN KEY(categoria_id) REFERENCES categorie(id),
            FOREIGN KEY(fornitore_id) REFERENCES fornitori(id),
            UNIQUE(fornitore_id, codice_fornitore)
        )""")

        # Capitoli con default %SG e %Utile
        cur.execute("""
        CREATE TABLE IF NOT EXISTS capitoli (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            codice TEXT NOT NULL UNIQUE,
            nome TEXT NOT NULL,
            cg_default_percentuale REAL DEFAULT 0.0,
            utile_default_percentuale REAL DEFAULT 0.0
        )""")

        # Voci di analisi
        cur.execute("""
        CREATE TABLE IF NOT EXISTS voci_analisi (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            capitolo_id INTEGER NOT NULL,
            codice TEXT NOT NULL,
            descrizione TEXT NOT NULL,
            costi_generali_percentuale REAL DEFAULT 0.0,
            utile_percentuale REAL DEFAULT 0.0,
            voce_unita_misura TEXT,
            voce_quantita REAL DEFAULT 1.0,
            FOREIGN KEY(capitolo_id) REFERENCES capitoli(id),
            UNIQUE(capitolo_id, codice)
        )""")

        # MIGRAZIONE: aggiunge la colonna prezzo_riferimento se manca
        try:
            cur.execute(
                "ALTER TABLE voci_analisi "
                "ADD COLUMN prezzo_riferimento REAL DEFAULT 0.0"
            )
        except sqlite3.OperationalError:
            pass  # già presente

        # Righe distinta
        cur.execute("""
        CREATE TABLE IF NOT EXISTS righe_distinta (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            voce_analisi_id INTEGER NOT NULL,
            materiale_id INTEGER NOT NULL,
            quantita REAL NOT NULL,
            FOREIGN KEY(voce_analisi_id) REFERENCES voci_analisi(id),
            FOREIGN KEY(materiale_id) REFERENCES materiali_base(id)
        )""")

        # Clienti / Preventivi
        cur.execute("""
        CREATE TABLE IF NOT EXISTS clienti (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nome TEXT NOT NULL,
            piva TEXT, indirizzo TEXT, cap TEXT, citta TEXT, provincia TEXT, nazione TEXT,
            email TEXT, telefono TEXT, note TEXT
        )""")
        cur.execute("""
        CREATE TABLE IF NOT EXISTS preventivi (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            numero TEXT NOT NULL,
            data TEXT NOT NULL,
            cliente_id INTEGER NOT NULL,
            note_finali TEXT,
            iva_percentuale REAL DEFAULT 22.0,
            imponibile REAL DEFAULT 0.0,
            iva_importo REAL DEFAULT 0.0,
            totale REAL DEFAULT 0.0,
            FOREIGN KEY(cliente_id) REFERENCES clienti(id)
        )""")
        cur.execute("""
        CREATE TABLE IF NOT EXISTS preventivo_righe (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            preventivo_id INTEGER NOT NULL,
            capitolo_id INTEGER NOT NULL,
            voce_id INTEGER NOT NULL,
            descrizione TEXT NOT NULL,
            note TEXT,
            um TEXT NOT NULL,
            quantita REAL NOT NULL,
            prezzo_unitario REAL NOT NULL,
            prezzo_totale REAL NOT NULL,
            FOREIGN KEY(preventivo_id) REFERENCES preventivi(id),
            FOREIGN KEY(capitolo_id) REFERENCES capitoli(id),
            FOREIGN KEY(voce_id) REFERENCES voci_analisi(id)
        )""")
        # Storico prezzi materiali
        cur.execute("""
        CREATE TABLE IF NOT EXISTS materiali_prezzi_storico (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            materiale_id INTEGER NOT NULL,
            prezzo_vecchio REAL NOT NULL,
            prezzo_nuovo REAL NOT NULL,
            changed_at TEXT NOT NULL DEFAULT (datetime('now')),
            note TEXT,
            FOREIGN KEY(materiale_id) REFERENCES materiali_base(id)
        )""")
        # Indici storico
        cur.execute("CREATE INDEX IF NOT EXISTS idx_sto_mat  ON materiali_prezzi_storico(materiale_id)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_sto_date ON materiali_prezzi_storico(changed_at)")
        # Trigger: logga i cambi prezzo dei materiali
        cur.execute("""
        CREATE TRIGGER IF NOT EXISTS trg_log_prezzo_materiale
        AFTER UPDATE OF prezzo_unitario ON materiali_base
        FOR EACH ROW
        WHEN NEW.prezzo_unitario IS NOT OLD.prezzo_unitario
        BEGIN
            INSERT INTO materiali_prezzi_storico (materiale_id, prezzo_vecchio, prezzo_nuovo, changed_at, note)
            VALUES (OLD.id, OLD.prezzo_unitario, NEW.prezzo_unitario, datetime('now'), 'Update da UI materiali');
        END;
        """)


        con.commit()

        # Seed iniziali
        if _exec(con, "SELECT COUNT(*) FROM categorie").fetchone()[0] == 0:
            _exec(con, "INSERT INTO categorie (nome) VALUES (?), (?), (?), (?)",
                  ["Edile", "Ferramenta", "Noleggi", "Pose"])
        if _exec(con, "SELECT COUNT(*) FROM fornitori").fetchone()[0] == 0:
            _exec(con, "INSERT INTO fornitori (nome) VALUES (?)", ["Fornitore Sconosciuto"])
        con.commit()

        # --- Indici utili ---
        cur.execute("CREATE INDEX IF NOT EXISTS idx_materiali_base_cat ON materiali_base(categoria_id)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_materiali_base_forn ON materiali_base(fornitore_id)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_voci_cap ON voci_analisi(capitolo_id)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_righe_voce ON righe_distinta(voce_analisi_id)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_righe_mat ON righe_distinta(materiale_id)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_prev_cliente ON preventivi(cliente_id)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_prev_data ON preventivi(data)")
        con.commit()

@contextmanager
def get_con():
    con = sqlite3.connect(DB_PATH)
    try:
        con.execute("PRAGMA foreign_keys = ON")  # FK attive su ogni connessione
        yield con
    finally:
        con.close()

# ------------------------------------------------------------------
# Query helpers
# ------------------------------------------------------------------
def df_categorie():
    with get_con() as con:
        return pd.read_sql_query("SELECT id, nome FROM categorie ORDER BY nome", con)

def df_fornitori():
    with get_con() as con:
        return pd.read_sql_query("""SELECT id, nome, piva, indirizzo, email, telefono
                                    FROM fornitori ORDER BY nome""", con)

def df_materiali():
    with get_con() as con:
        return pd.read_sql_query("""
            SELECT m.id,
                   m.categoria_id, c.nome AS categoria,
                   m.fornitore_id, f.nome AS fornitore,
                   m.codice_fornitore, m.descrizione, m.unita_misura,
                   IFNULL(m.quantita_default,1.0) AS quantita_default,
                   m.prezzo_unitario
            FROM materiali_base m
            JOIN categorie c  ON c.id = m.categoria_id
            JOIN fornitori f  ON f.id = m.fornitore_id
            ORDER BY c.nome, f.nome, m.codice_fornitore
        """, con)

def df_capitoli():
    with get_con() as con:
        return pd.read_sql_query("""
            SELECT id, codice, nome,
                   IFNULL(cg_default_percentuale,0) AS cg_def,
                   IFNULL(utile_default_percentuale,0) AS ut_def
            FROM capitoli ORDER BY codice
        """, con)

def df_voci(capitolo_id: Optional[int] = None):
    with get_con() as con:
        if capitolo_id:
            q = """
            SELECT v.id, v.capitolo_id, c.codice AS capitolo_codice, c.nome AS capitolo_nome,
                   v.codice, v.descrizione,
                   IFNULL(v.costi_generali_percentuale,0) AS cg_pct,
                   IFNULL(v.utile_percentuale,0) AS utile_pct,
                   v.voce_unita_misura AS um_voce,
                   IFNULL(v.voce_quantita,1.0) AS q_voce,
                   IFNULL(v.prezzo_riferimento,0.0) AS prezzo_rif
            FROM voci_analisi v
            JOIN capitoli c ON c.id = v.capitolo_id
            WHERE v.capitolo_id = ?
            ORDER BY c.codice, v.codice
            """
            return pd.read_sql_query(q, con, params=[capitolo_id])
        else:
            q = """
            SELECT v.id, v.capitolo_id, c.codice AS capitolo_codice, c.nome AS capitolo_nome,
                   v.codice, v.descrizione,
                   IFNULL(v.costi_generali_percentuale,0) AS cg_pct,
                   IFNULL(v.utile_percentuale,0) AS utile_pct,
                   v.voce_unita_misura AS um_voce,
                   IFNULL(v.voce_quantita,1.0) AS q_voce,
                   IFNULL(v.prezzo_riferimento,0.0) AS prezzo_rif
            FROM voci_analisi v
            JOIN capitoli c ON c.id = v.capitolo_id
            ORDER BY c.codice, v.codice
            """
            return pd.read_sql_query(q, con)

def df_righe(voce_id: int):
    with get_con() as con:
        return pd.read_sql_query("""
            SELECT r.id, r.voce_analisi_id, r.materiale_id, r.quantita,
                   m.descrizione AS materiale_descrizione,
                   m.unita_misura, m.prezzo_unitario,
                   c.nome AS categoria, f.nome AS fornitore, m.codice_fornitore,
                   (r.quantita * m.prezzo_unitario) AS subtotale
            FROM righe_distinta r
            JOIN materiali_base m ON m.id = r.materiale_id
            JOIN categorie c ON c.id = m.categoria_id
            JOIN fornitori f ON f.id = m.fornitore_id
            WHERE r.voce_analisi_id = ?
            ORDER BY r.id
        """, con, params=[voce_id])

def get_voce(voce_id: int) -> Optional[dict]:
    with get_con() as con:
        row = _exec(con, """
            SELECT v.id, v.capitolo_id, c.codice, c.nome,
                   v.codice, v.descrizione,
                   IFNULL(v.costi_generali_percentuale,0),
                   IFNULL(v.utile_percentuale,0),
                   v.voce_unita_misura,
                   IFNULL(v.voce_quantita,1.0),
                   IFNULL(v.prezzo_riferimento,0.0)
            FROM voci_analisi v
            JOIN capitoli c ON c.id = v.capitolo_id
            WHERE v.id = ?
        """, (voce_id,)).fetchone()
        if not row:
            return None
        return {
            "id": row[0], "capitolo_id": row[1],
            "capitolo_codice": row[2], "capitolo_nome": row[3],
            "codice": row[4], "descrizione": row[5],
            "cg_pct": float(row[6]), "utile_pct": float(row[7]),
            "um_voce": row[8], "q_voce": float(row[9]),
            "prezzo_rif": float(row[10]),
        }

# ------------------------------------------------------------------
# Calcoli
# ------------------------------------------------------------------
def compute_totali_voce(voce_id: int) -> Dict[str, float]:
    df = df_righe(voce_id)
    costo_materie = float(df["subtotale"].sum()) if not df.empty else 0.0
    voce = get_voce(voce_id) or {"cg_pct": 0.0, "utile_pct": 0.0}
    cg = costo_materie * (voce["cg_pct"] / 100.0)
    base = costo_materie + cg
    utile = base * (voce["utile_pct"] / 100.0)
    totale = base + utile
    return {
        "costo_materie": costo_materie,
        "costi_generali": cg,
        "utile": utile,
        "cg_pct": voce["cg_pct"],
        "utile_pct": voce["utile_pct"],
        "totale": totale,
    }
# -------- (2) Impatti da aggiornamento materiali --------
def voci_impattate_da_materiali(material_ids: list[int]) -> pd.DataFrame:
    """Ritorna le voci che usano almeno uno dei materiali indicati."""
    if not material_ids:
        return pd.DataFrame(columns=["voce_id","capitolo_codice","capitolo_nome","codice","descrizione","prezzo_rif"])
    with get_con() as con:
        q = """
        SELECT DISTINCT v.id AS voce_id,
               c.codice AS capitolo_codice, c.nome AS capitolo_nome,
               v.codice, v.descrizione,
               IFNULL(v.prezzo_riferimento,0.0) AS prezzo_rif
        FROM righe_distinta r
        JOIN voci_analisi v ON v.id = r.voce_analisi_id
        JOIN capitoli c    ON c.id = v.capitolo_id
        WHERE r.materiale_id IN ({})
        ORDER BY c.codice, v.codice
        """.format(",".join(["?"]*len(material_ids)))
        return pd.read_sql_query(q, con, params=list(material_ids))

def anteprima_impatti_materiali(material_ids: list[int]) -> pd.DataFrame:
    """
    Calcola il totale attuale della VOCE (con i prezzi base correnti) e lo
    confronta con il prezzo di riferimento della voce (se presente).
    """
    voci_df = voci_impattate_da_materiali(material_ids)
    rows = []
    for _, r in voci_df.iterrows():
        tot = compute_totali_voce(int(r.voce_id))["totale"]
        rif = float(r.get("prezzo_rif", 0.0))
        delta_pct = ((tot - rif) / rif * 100.0) if rif > 0 else None
        rows.append({
            "Capitolo": r.capitolo_codice,
            "Voce": r.codice,
            "Descrizione": r.descrizione,
            "Totale attuale (€)": round(tot, 2),
            "Prezzo riferimento (€)": (round(rif, 2) if rif > 0 else "-"),
            "Δ vs riferimento (%)": (f"{delta_pct:+.2f}%" if delta_pct is not None else "-"),
            "voce_id": int(r.voce_id),
        })
    return pd.DataFrame(rows)

def prezzo_unitario_voce(voce_id: int) -> float:
    v = get_voce(voce_id)
    if not v:
        return 0.0
    tot = compute_totali_voce(voce_id)["totale"]
    q = max(float(v["q_voce"]), 1e-9)
    return tot / q

# ------------------------------------------------------------------
# Mutations (CRUD)
# ------------------------------------------------------------------
def add_categoria(nome: str):
    with get_con() as con:
        try:
            _exec(con, "INSERT INTO categorie (nome) VALUES (?)", (nome.strip(),))
            con.commit()
            st.success("Categoria aggiunta.")
        except sqlite3.IntegrityError:
            st.warning("Categoria già esistente.")

def delete_categoria(cid: int):
    with get_con() as con:
        used = _exec(con, "SELECT COUNT(*) FROM materiali_base WHERE categoria_id=?", (cid,)).fetchone()[0]
        if used:
            st.warning("Impossibile eliminare: categoria usata da materiali.")
            return
        _exec(con, "DELETE FROM categorie WHERE id=?", (cid,))
        con.commit()
        st.success("Categoria eliminata.")

def add_fornitore(nome, piva, indirizzo, email, telefono):
    """Inserisce un fornitore solo se NON esiste già per Nome (normalizzato) o P.IVA (solo cifre)."""
    nome_n = _norm_text(nome)
    piva_n = _digits_only(piva)

    with get_con() as con:
        # Controllo duplicati lato applicativo (robusto contro varianti di spazi/maiuscole/punteggiatura)
        rows = _exec(con, "SELECT id, nome, piva FROM fornitori").fetchall()
        for fid, fn, fp in rows:
            if _norm_text(fn) == nome_n:
                st.error("Fornitore già esistente: il NOME coincide. Operazione annullata.")
                return
            if piva_n and _digits_only(fp) == piva_n:
                st.error("Fornitore già esistente: la P.IVA coincide. Operazione annullata.")
                return

        # Inserimento (gestisce anche l'UNIQUE(nome) a schema)
        try:
            _exec(con, """INSERT INTO fornitori (nome,piva,indirizzo,email,telefono)
                          VALUES (?,?,?,?,?)""",
                  (nome.strip(), piva, indirizzo, email, telefono))
            con.commit()
            st.success("Fornitore aggiunto.")
        except sqlite3.IntegrityError:
            st.warning("Fornitore già esistente (vincolo su Nome).")

def delete_fornitore(fid: int):
    with get_con() as con:
        used = _exec(con, "SELECT COUNT(*) FROM materiali_base WHERE fornitore_id=?", (fid,)).fetchone()[0]
        if used:
            st.warning("Impossibile eliminare: fornitore usato da materiali.")
            return
        _exec(con, "DELETE FROM fornitori WHERE id=?", (fid,))
        con.commit()
        st.success("Fornitore eliminato.")

def add_materiale(categoria_id, fornitore_id, codice_fornitore, descrizione, um, qdef, prezzo):
    try:
        with get_con() as con:
            _exec(con, """
                INSERT INTO materiali_base (categoria_id, fornitore_id, codice_fornitore, descrizione, unita_misura, quantita_default, prezzo_unitario)
                VALUES (?, ?, ?, ?, ?, ?, ?)
            """, (int(categoria_id), int(fornitore_id), codice_fornitore.strip(), descrizione.strip(),
                  um, float(qdef or 1.0), float(prezzo)))
            con.commit()
            st.success("Materiale inserito.")
    except sqlite3.IntegrityError:
        st.error("Codice fornitore già presente per questo fornitore.")

def update_materiali_bulk(df_edit: pd.DataFrame, df_orig: pd.DataFrame):
    changes = []
    for _, row in df_edit.iterrows():
        orig = df_orig[df_orig["id"] == row["id"]].iloc[0]
        fields = ["descrizione", "unita_misura", "quantita_default", "prezzo_unitario"]
        updates = {f: row[f] for f in fields if str(row[f]) != str(orig[f])}
        if updates:
            changes.append((int(row["id"]), updates))
    if not changes:
        st.info("Nessuna modifica da salvare.")
        return
    with get_con() as con:
        for mid, upd in changes:
            sets = ", ".join([f"{k}=?" for k in upd.keys()])
            vals = list(upd.values()) + [mid]
            _exec(con, f"UPDATE materiali_base SET {sets} WHERE id=?", vals)
        con.commit()
    st.success(f"Salvate {len(changes)} modifiche.")

def add_capitolo(codice, nome, cg_def, ut_def):
    try:
        with get_con() as con:
            _exec(con, "INSERT INTO capitoli (codice, nome, cg_default_percentuale, utile_default_percentuale) VALUES (?,?,?,?)",
                  (codice.strip(), nome.strip(), float(cg_def or 0.0), float(ut_def or 0.0)))
            con.commit()
            st.success("Capitolo inserito.")
    except sqlite3.IntegrityError:
        st.error("Codice capitolo già esistente.")

def update_capitolo_defaults(cid: int, cg_def: float, ut_def: float):
    with get_con() as con:
        _exec(con, "UPDATE capitoli SET cg_default_percentuale=?, utile_default_percentuale=? WHERE id=?",
              (float(cg_def or 0.0), float(ut_def or 0.0), int(cid)))
        con.commit()
        st.success("Aggiornati i valori di Spese generali e Utile per il capitolo (influenza nuove voci; le esistenti restano invariate).")

def delete_capitolo(cid: int):
    with get_con() as con:
        used = _exec(con, "SELECT COUNT(*) FROM voci_analisi WHERE capitolo_id=?", (cid,)).fetchone()[0]
        if used:
            st.warning("Impossibile eliminare: il capitolo contiene voci.")
            return
        _exec(con, "DELETE FROM capitoli WHERE id=?", (cid,))
        con.commit()
        st.success("Capitolo eliminato.")

def add_voce(capitolo_id, codice, descrizione, cg_pct, utile_pct, um_voce, q_voce, prezzo_rif=0.0):
    try:
        with get_con() as con:
            _exec(con, """INSERT INTO voci_analisi
                          (capitolo_id, codice, descrizione, costi_generali_percentuale, utile_percentuale,
                           voce_unita_misura, voce_quantita, prezzo_riferimento)
                          VALUES (?,?,?,?,?,?,?,?)""",
                  (int(capitolo_id), codice.strip(), descrizione.strip(), float(cg_pct or 0.0),
                   float(utile_pct or 0.0), um_voce, float(q_voce or 1.0), float(prezzo_rif or 0.0)))
            con.commit()
            st.success("Voce creata.")
    except sqlite3.IntegrityError:
        st.error("Codice voce già esistente nel capitolo.")

def update_voce_perc(vid: int, cg_pct: float, utile_pct: float):
    with get_con() as con:
        _exec(con, "UPDATE voci_analisi SET costi_generali_percentuale=?, utile_percentuale=? WHERE id=?",
              (float(cg_pct or 0.0), float(utile_pct or 0.0), int(vid)))
        con.commit()
        st.success("Percentuali aggiornate.")

def update_voce_perc_umqty(vid: int, cg_pct: float, utile_pct: float, um_voce: str, q_voce: float, prezzo_rif: float):
    with get_con() as con:
        _exec(con, """UPDATE voci_analisi
                      SET costi_generali_percentuale=?, utile_percentuale=?, voce_unita_misura=?, voce_quantita=?, prezzo_riferimento=?
                      WHERE id=?""",
              (float(cg_pct or 0.0), float(utile_pct or 0.0), um_voce, float(q_voce or 1.0),
               float(prezzo_rif or 0.0), int(vid)))
        con.commit()
        st.success("Voce aggiornata.")

def add_riga_distinta(voce_id: int, materiale_id: int, quantita: float):
    with get_con() as con:
        _exec(con, "INSERT INTO righe_distinta (voce_analisi_id, materiale_id, quantita) VALUES (?,?,?)",
              (int(voce_id), int(materiale_id), float(quantita)))
        con.commit()
        st.success("Riga aggiunta.")

def update_quantita_righe(voce_id: int, edited: pd.DataFrame, original: pd.DataFrame):
    diffs = []
    for _, r in edited.iterrows():
        o = original[original["id"] == r["id"]].iloc[0]
        if float(r["quantita"]) != float(o["quantita"]):
            diffs.append((float(r["quantita"]), int(r["id"])))
    if not diffs:
        st.info("Nessuna quantità modificata.")
        return
    with get_con() as con:
        for q, rid in diffs:
            _exec(con, "UPDATE righe_distinta SET quantita=? WHERE id=?", (q, rid))
        con.commit()
    st.success(f"Aggiornate {len(diffs)} righe.")

def delete_riga(riga_id: int):
    with get_con() as con:
        _exec(con, "DELETE FROM righe_distinta WHERE id=?", (riga_id,))
        con.commit()
        st.success("Riga eliminata.")

def delete_voce(vid: int):
    with get_con() as con:
        _exec(con, "DELETE FROM righe_distinta WHERE voce_analisi_id=?", (vid,))
        _exec(con, "DELETE FROM voci_analisi WHERE id=?", (vid,))
        con.commit()
        st.success("Voce eliminata.")

def clone_voce(vid: int):
    v = get_voce(vid)
    if not v:
        st.error("Voce non trovata.")
        return
    new_code = f"{v['codice']}-COPY"
    with get_con() as con:
        try:
            _exec(con, """INSERT INTO voci_analisi (capitolo_id, codice, descrizione, costi_generali_percentuale, utile_percentuale, voce_unita_misura, voce_quantita, prezzo_riferimento)
                          VALUES (?,?,?,?,?,?,?,?)""",
                  (v["capitolo_id"], new_code, v["descrizione"], v["cg_pct"], v["utile_pct"], v["um_voce"], v["q_voce"], v.get("prezzo_rif", 0.0)))
            new_id = _exec(con, "SELECT last_insert_rowid()").fetchone()[0]
            rows = _exec(con, "SELECT materiale_id, quantita FROM righe_distinta WHERE voce_analisi_id=?", (vid,)).fetchall()
            for m_id, q in rows:
                _exec(con, "INSERT INTO righe_distinta (voce_analisi_id, materiale_id, quantita) VALUES (?,?,?)",
                      (new_id, m_id, q))
            con.commit()
            st.success(f"Voce clonata come codice {new_code}.")
        except sqlite3.IntegrityError:
            st.error("Esiste già una voce con quel codice; riprova.")

# --- Eliminazioni con controlli di collegamenti ---
def delete_cliente(cid: int):
    with get_con() as con:
        used = _exec(con, "SELECT COUNT(*) FROM preventivi WHERE cliente_id=?", (cid,)).fetchone()[0]
        if used:
            st.warning("Impossibile eliminare: il cliente ha preventivi collegati.")
            return
        _exec(con, "DELETE FROM clienti WHERE id=?", (cid,))
        con.commit()
        st.success("Cliente eliminato.")

def delete_materiale(mid: int):
    with get_con() as con:
        used = _exec(con, "SELECT COUNT(*) FROM righe_distinta WHERE materiale_id=?", (mid,)).fetchone()[0]
        if used:
            st.warning("Impossibile eliminare: materiale presente in almeno una voce di analisi.")
            return
        _exec(con, "DELETE FROM materiali_base WHERE id=?", (mid,))
        con.commit()
        st.success("Materiale eliminato.")

# ------------------------------------------------------------------
# Import/Export
# ------------------------------------------------------------------
def import_materiali_csv(file):
    try:
        df = pd.read_csv(file)
    except Exception:
        file.seek(0)
        df = pd.read_excel(file)

    required = {"categoria","fornitore","codice_fornitore","descrizione","unita_misura","prezzo_unitario"}
    cols_lower = {c.lower(): c for c in df.columns}
    if not required.issubset(set(cols_lower.keys())):
        missing = required - set(cols_lower.keys())
        st.error(f"Colonne mancanti: {', '.join(missing)}")
        return

    df = df.rename(columns={v: k.lower() for k, v in cols_lower.items()})
    if "quantita_default" not in df.columns:
        df["quantita_default"] = 1.0

    with get_con() as con:
        for _, r in df.iterrows():
            um = str(r["unita_misura"]).strip()
            if um not in UM_CHOICES:
                st.warning(f"UM non valida '{um}' per codice {r['codice_fornitore']} → saltato.")
                continue

            cat = str(r["categoria"]).strip()
            row = _exec(con, "SELECT id FROM categorie WHERE nome=?", (cat,)).fetchone()
            if not row:
                _exec(con, "INSERT INTO categorie (nome) VALUES (?)", (cat,))
                cat_id = _exec(con, "SELECT last_insert_rowid()").fetchone()[0]
            else:
                cat_id = row[0]

            forn = str(r["fornitore"]).strip()
            row = _exec(con, "SELECT id FROM fornitori WHERE nome=?", (forn,)).fetchone()
            if not row:
                _exec(con, "INSERT INTO fornitori (nome) VALUES (?)", (forn,))
                forn_id = _exec(con, "SELECT last_insert_rowid()").fetchone()[0]
            else:
                forn_id = row[0]

            try:
                _exec(con, """INSERT INTO materiali_base
                              (categoria_id, fornitore_id, codice_fornitore, descrizione, unita_misura, quantita_default, prezzo_unitario)
                              VALUES (?,?,?,?,?,?,?)""",
                      (cat_id, forn_id,
                       str(r["codice_fornitore"]).strip(),
                       str(r["descrizione"]).strip(),
                       um,
                       _to_float(r.get("quantita_default", 1.0), 1.0),
                       _to_float(r["prezzo_unitario"], 0.0)))
                con.commit()
            except sqlite3.IntegrityError:
                st.warning(f"Duplicato: {forn} / {r['codice_fornitore']} → saltato.")
    st.success("Import materiali completato.")

def import_fornitori_csv(file):
    try:
        df = pd.read_csv(file)
    except Exception:
        file.seek(0)
        df = pd.read_excel(file)

    cols = {c.lower(): c for c in df.columns}
    if "nome" not in cols:
        st.error("Colonna obbligatoria mancante: 'nome'")
        return

    df = df.rename(columns={v: k.lower() for k, v in cols.items()})
    with get_con() as con:
        inserted, skipped = 0, 0
        for _, r in df.iterrows():
            name = str(r["nome"]).strip()
            if not name:
                skipped += 1
                continue
            exists = _exec(con, "SELECT 1 FROM fornitori WHERE nome=?", (name,)).fetchone()
            if exists:
                skipped += 1
                continue
            _exec(con, """INSERT INTO fornitori (nome,piva,indirizzo,email,telefono)
                          VALUES (?,?,?,?,?)""",
                  (name, str(r.get("piva") or ""), str(r.get("indirizzo") or ""),
                   str(r.get("email") or ""), str(r.get("telefono") or "")))
            inserted += 1
        con.commit()
    st.success(f"Import fornitori completato. Inseriti: {inserted}, saltati: {skipped}.")

def export_excel():  # SOLO Sommario EPU con Nome Capitolo (come richiesto)
    voci = df_voci()
    if voci.empty:
        st.warning("Non ci sono voci da esportare.")
        return None

    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        rows = []
        for _, r in voci.iterrows():
            tot = compute_totali_voce(int(r.id))
            rows.append({
                "Codice Capitolo": r.capitolo_codice,
                "Nome Capitolo": r.capitolo_nome,
                "Cod. Voce": r.codice,
                "Descrizione Voce": r.descrizione,
                "UM Voce": r.um_voce,
                "Q.tà Voce": r.q_voce,
                "CG %": r.cg_pct,
                "Utile %": r.utile_pct,
                "Materie (€)": round(tot["costo_materie"], 2),
                "Spese generali (€)": round(tot["costi_generali"], 2),
                "Utile (€)": round(tot["utile"], 2),
                "Totale (€)": round(tot["totale"], 2),
            })
        df_sommario = pd.DataFrame(rows).sort_values(["Codice Capitolo", "Cod. Voce"])
        df_sommario.to_excel(writer, index=False, sheet_name="Sommario EPU")
    buffer.seek(0)
    return buffer

# ------------------------------------------------------------------
# CLIENTI / PREVENTIVI
# ------------------------------------------------------------------
def df_clienti():
    with get_con() as con:
        return pd.read_sql_query("""
            SELECT id, nome, piva, indirizzo, cap, citta, provincia, nazione, email, telefono, note
            FROM clienti ORDER BY nome
        """, con)

def add_cliente(**kwargs):
    with get_con() as con:
        _exec(con, """INSERT INTO clienti (nome,piva,indirizzo,cap,citta,provincia,nazione,email,telefono,note)
                      VALUES (?,?,?,?,?,?,?,?,?,?)""",
              (kwargs.get("nome","").strip(), kwargs.get("piva",""), kwargs.get("indirizzo",""),
               kwargs.get("cap",""), kwargs.get("citta",""), kwargs.get("provincia",""), kwargs.get("nazione",""),
               kwargs.get("email",""), kwargs.get("telefono",""), kwargs.get("note","")))
        con.commit()
        st.success("Cliente inserito.")

def create_preventivo(numero: str, data_iso: str, cliente_id: int, note_finali: str, iva_percent: float) -> int:
    with get_con() as con:
        _exec(con, """INSERT INTO preventivi (numero,data,cliente_id,note_finali,iva_percentuale,imponibile,iva_importo,totale)
                      VALUES (?,?,?,?,?,?,?,?)""",
              (numero.strip(), data_iso, int(cliente_id), note_finali, float(iva_percent or 0.0), 0.0, 0.0, 0.0))
        pid = _exec(con, "SELECT last_insert_rowid()").fetchone()[0]
        con.commit()
        return int(pid)

def add_riga_preventivo(pid: int, capitolo_id: int, voce_id: int, descrizione: str, note: str,
                        um: str, quantita: float, prezzo_unitario: float):
    prezzo_totale = float(quantita) * float(prezzo_unitario)
    with get_con() as con:
        _exec(con, """INSERT INTO preventivo_righe
                      (preventivo_id,capitolo_id,voce_id,descrizione,note,um,quantita,prezzo_unitario,prezzo_totale)
                      VALUES (?,?,?,?,?,?,?,?,?)""",
              (int(pid), int(capitolo_id), int(voce_id), descrizione.strip(), note, um, float(quantita),
               float(prezzo_unitario), prezzo_totale))
        con.commit()

def df_preventivo(pid: int):
    with get_con() as con:
        testa = pd.read_sql_query("""
            SELECT p.id, p.numero, p.data, p.cliente_id, p.note_finali, p.iva_percentuale, p.imponibile, p.iva_importo, p.totale,
                   c.nome AS cliente_nome, c.piva, c.indirizzo, c.cap, c.citta, c.provincia, c.nazione, c.email, c.telefono
            FROM preventivi p
            JOIN clienti c ON c.id = p.cliente_id
            WHERE p.id = ?
        """, con, params=[pid])
        righe = pd.read_sql_query("""
            SELECT r.id, r.preventivo_id, r.capitolo_id, cap.codice AS capitolo_codice, cap.nome AS capitolo_nome,
                   r.voce_id, v.codice AS voce_codice,
                   r.descrizione, r.note, r.um, r.quantita, r.prezzo_unitario, r.prezzo_totale
            FROM preventivo_righe r
            JOIN capitoli cap ON cap.id = r.capitolo_id
            JOIN voci_analisi v ON v.id = r.voce_id
            WHERE r.preventivo_id = ?
            ORDER BY cap.codice, v.codice, r.id
        """, con, params=[pid])
        return testa, righe

def ricalcola_totali_preventivo(pid: int, iva_percent: Optional[float] = None):
    with get_con() as con:
        imp = _exec(con, "SELECT IFNULL(SUM(prezzo_totale),0) FROM preventivo_righe WHERE preventivo_id=?", (pid,)).fetchone()[0]
        if iva_percent is None:
            iva_percent = _exec(con, "SELECT iva_percentuale FROM preventivi WHERE id=?", (pid,)).fetchone()[0]
        iva_imp = imp * float(iva_percent) / 100.0
        tot = imp + iva_imp
        _exec(con, "UPDATE preventivi SET imponibile=?, iva_percentuale=?, iva_importo=?, totale=? WHERE id=?",
              (imp, float(iva_percent), iva_imp, tot, pid))
        con.commit()

def export_preventivo_excel(pid: int):
    testa, righe = df_preventivo(pid)
    if testa.empty:
        return None
    buffer = io.BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as w:
        testa.to_excel(w, index=False, sheet_name="Testata")
        righe.to_excel(w, index=False, sheet_name="Righe")
        if not righe.empty:
            bycap = (righe.groupby(["capitolo_codice","capitolo_nome"])["prezzo_totale"]
                     .sum().reset_index().rename(columns={"prezzo_totale":"Totale capitolo (€)"}))
        else:
            bycap = pd.DataFrame(columns=["capitolo_codice","capitolo_nome","Totale capitolo (€)"])
        bycap.to_excel(w, index=False, sheet_name="Totali capitoli")
        riepilogo = testa[["numero","data","cliente_nome","imponibile","iva_percentuale","iva_importo","totale"]].copy()
        riepilogo = riepilogo.rename(columns={"numero":"Numero","data":"Data","cliente_nome":"Cliente",
                                              "imponibile":"Imponibile (€)","iva_percentuale":"IVA %","iva_importo":"IVA (€)","totale":"Totale (€)"})
        riepilogo.to_excel(w, index=False, sheet_name="Riepilogo")
    buffer.seek(0)
    return buffer

def df_preventivi_archivio(numero_like: str = "", data_like: str = "", cliente_id: Optional[int] = None):
    with get_con() as con:
        q = """
        SELECT p.id, p.numero, p.data,
               COALESCE(c.nome, '[cliente mancante]') AS cliente,
               p.imponibile, p.iva_percentuale, p.totale
        FROM preventivi p
        LEFT JOIN clienti c ON c.id = p.cliente_id
        WHERE 1=1
        """
        params = []
        if numero_like:
            q += " AND p.numero LIKE ?"; params.append(f"%{numero_like}%")
        if data_like:
            q += " AND p.data LIKE ?"; params.append(f"%{data_like}%")
        if cliente_id:
            q += " AND p.cliente_id = ?"; params.append(int(cliente_id))
        q += " ORDER BY p.data DESC, p.numero DESC"
        return pd.read_sql_query(q, con, params=params)

def export_preventivo_docx(pid: int):
    from docx import Document
    from docx.shared import Pt, Cm

    testa, righe = df_preventivo(pid)
    if testa.empty:
        return None

    d = Document()
    d.add_heading(f"Preventivo {testa['numero'].iloc[0]} del {testa['data'].iloc[0]}", level=1)

    # Cliente
    cli = [
        f"Cliente: {testa['cliente_nome'].iloc[0]}",
        f"P.IVA/CF: {testa['piva'].iloc[0] or '-'}",
        f"Indirizzo: {testa['indirizzo'].iloc[0] or '-'}",
        f"Città: {testa['cap'].iloc[0] or ''} {testa['citta'].iloc[0] or ''} ({testa['provincia'].iloc[0] or ''})",
        f"Nazione: {testa['nazione'].iloc[0] or '-'}",
        f"Email: {testa['email'].iloc[0] or '-'}  Tel: {testa['telefono'].iloc[0] or '-'}",
    ]
    for r in cli:
        d.add_paragraph(r)

    d.add_paragraph("")  # spazio

    # Tabella righe
    table = d.add_table(rows=1, cols=7)
    hdr = table.rows[0].cells
    hdr[0].text = "Capitolo"
    hdr[1].text = "Voce"
    hdr[2].text = "Descrizione"
    hdr[3].text = "UM"
    hdr[4].text = "Q.tà"
    hdr[5].text = "Prezzo U (€)"
    hdr[6].text = "Totale (€)"

    for _, r in righe.iterrows():
        row = table.add_row().cells
        row[0].text = f"{r['capitolo_codice']} {r['capitolo_nome']}"
        row[1].text = str(r["voce_codice"])
        row[2].text = str(r["descrizione"])
        row[3].text = str(r["um"])
        row[4].text = f"{r['quantita']:.2f}"
        row[5].text = f"{r['prezzo_unitario']:.2f}"
        row[6].text = f"{r['prezzo_totale']:.2f}"

        note_val = r["note"] if "note" in r and pd.notna(r["note"]) and str(r["note"]).strip() else ""
        if note_val:
            d.add_paragraph(f"Note: {note_val}")

    d.add_paragraph("")

    # Totali per capitolo
    if not righe.empty:
        bycap = (righe.groupby(["capitolo_codice","capitolo_nome"])["prezzo_totale"].sum()
                 .reset_index().rename(columns={"prezzo_totale":"Totale capitolo (€)"}))
        d.add_paragraph("Totali per capitolo:")
        for _, rr in bycap.iterrows():
            d.add_paragraph(f"- {rr['capitolo_codice']} {rr['capitolo_nome']}: € {rr['Totale capitolo (€)']:.2f}")

    d.add_paragraph("")

    # Riepilogo documento
    imp = float(testa["imponibile"].iloc[0])
    iva_p = float(testa["iva_percentuale"].iloc[0])
    iva_imp = float(testa["iva_importo"].iloc[0])
    tot = float(testa["totale"].iloc[0])

    d.add_paragraph(f"Imponibile: € {imp:.2f}")
    d.add_paragraph(f"IVA {iva_p:.0f}%: € {iva_imp:.2f}")
    d.add_paragraph(f"Totale documento: € {tot:.2f}")

    if testa["note_finali"].iloc[0]:
        d.add_paragraph("")
        d.add_paragraph(f"Note finali: {testa['note_finali'].iloc[0]}")

    buf = io.BytesIO()
    d.save(buf)
    buf.seek(0)
    return buf

# ------------------------------------------------------------------
# UI – Categorie
# ------------------------------------------------------------------
def ui_categorie():
    st.subheader("Categorie")
    df = df_categorie()
    left, right = st.columns([2, 1])
    with left:
        st.dataframe(df, use_container_width=True, hide_index=True)
    with right:
        nome = st.text_input("Nuova categoria")
        if st.button("➕ Aggiungi categoria") and nome:
            add_categoria(nome); st.rerun()
        if not df.empty:
            del_id = st.selectbox("Elimina categoria", options=[None]+df["id"].tolist(),
                                  format_func=lambda x: "—" if x is None else df[df["id"]==x]["nome"].iloc[0])
            if del_id and st.button("Elimina"):
                delete_categoria(int(del_id)); st.rerun()

# ------------------------------------------------------------------
# UI – Fornitori
# ------------------------------------------------------------------
def ui_fornitori():
    st.subheader("Fornitori")
    df = df_fornitori()
    st.dataframe(df, use_container_width=True, hide_index=True, height=260)
    with st.expander("➕ Nuovo fornitore"):
        c1, c2 = st.columns(2)
        nome = c1.text_input("Nome *")
        piva = c1.text_input("P.IVA")
        indirizzo = c2.text_input("Indirizzo")
        email = c2.text_input("Email")
        telefono = c1.text_input("Telefono", key="tel_forn")
        if st.button("Aggiungi") and nome:
            add_fornitore(nome, piva, indirizzo, email, telefono); st.rerun()

    with st.expander("📥 Import fornitori da CSV/Excel"):
        st.markdown("**Schema**: `nome` (obbl.), `piva`, `indirizzo`, `email`, `telefono` (opzionali).")
        up = st.file_uploader("Carica .csv o .xlsx", type=["csv","xlsx"], key="up_fornitori")
        if up is not None:
            import_fornitori_csv(up); st.rerun()

    if not df.empty:
        fid = st.selectbox("Elimina fornitore", options=[None]+df["id"].tolist(),
                           format_func=lambda x: "—" if x is None else df[df["id"]==x]["nome"].iloc[0])
        if fid and st.button("Elimina fornitore"):
            delete_fornitore(int(fid)); st.rerun()

# ------------------------------------------------------------------
# UI – Materiali (con filtri stile Excel)
# ------------------------------------------------------------------
def ui_materiali():
    st.subheader("Archivio prezzi base (Materiali)")

    cat_df = df_categorie(); forn_df = df_fornitori()
    if cat_df.empty or forn_df.empty:
        st.info("Servono almeno 1 Categoria e 1 Fornitore.")
        return

    # ---------------------------
    # Form inserimento materiale
    # ---------------------------
    with st.form("form_materiale"):
        c1, c2, c3, c4 = st.columns([1.2, 1.2, 1.2, 1.6])
        categoria_id = c1.selectbox(
            "Categoria",
            options=cat_df["id"],
            format_func=lambda i: cat_df[cat_df["id"] == i]["nome"].iloc[0],
            key="mat_cat"
        )
        fornitore_id = c2.selectbox(
            "Fornitore",
            options=forn_df["id"],
            format_func=lambda i: forn_df[forn_df["id"] == i]["nome"].iloc[0],
            key="mat_forn"
        )
        codice_fornitore = c3.text_input("Codice fornitore *", key="mat_cod")
        um = c4.selectbox("UM", UM_CHOICES, key="mat_um")

        descrizione = st.text_input("Descrizione *", key="mat_desc")
        colq, colp = st.columns(2)
        qdef = colq.number_input("Quantità default", min_value=0.0, value=1.0, step=1.0, key="mat_qdef")
        prezzo = colp.number_input("Prezzo unitario (€) *", min_value=0.0, value=0.0, step=0.01, format="%.2f", key="mat_price")

        if st.form_submit_button("➕ Aggiungi materiale"):
            if not (codice_fornitore and descrizione and prezzo > 0):
                st.warning("Compila i campi obbligatori contrassegnati con *.")
            else:
                add_materiale(categoria_id, fornitore_id, codice_fornitore, descrizione, um, qdef, prezzo)
                st.rerun()

    # ---------------------------
    # Tabella + filtri stile Excel
    # ---------------------------
    df_all = df_materiali()
    if df_all is None or df_all.empty:
        st.info("Nessun materiale in archivio.")
        return

    # Filtri per colonna (in testa)
    fc1, fc2, fc3, fc4 = st.columns([1, 1, 1, 2])
    f_categoria = fc1.text_input("Filtro Categoria", key="flt_mat_categoria")
    f_fornitore = fc2.text_input("Filtro Fornitore", key="flt_mat_fornitore")
    f_codforn   = fc3.text_input("Filtro Cod. Fornitore", key="flt_mat_codforn")
    f_descr     = fc4.text_input("Filtro Descrizione", key="flt_mat_descr")

    dfv = df_all.copy()
    mask = (
        like_mask(dfv["categoria"], f_categoria)
        & like_mask(dfv["fornitore"], f_fornitore)
        & like_mask(dfv["codice_fornitore"], f_codforn)
        & like_mask(dfv["descrizione"], f_descr)
    )
    dfv = dfv[mask]

    st.caption(f"{len(dfv)} materiali (totale archivio: {len(df_all)})")

    # Vista + editor (solo alcune colonne modificabili)
    view_cols = ["id","categoria","fornitore","codice_fornitore","descrizione","unita_misura","quantita_default","prezzo_unitario"]
    view = dfv[view_cols].copy()

    edited = st.data_editor(
        view,
        use_container_width=True,
        hide_index=True,
        num_rows="fixed",
        column_config={
            "unita_misura": st.column_config.SelectboxColumn("UM", options=UM_CHOICES),
            "quantita_default": st.column_config.NumberColumn("quantita_default", step=0.1, min_value=0.0),
            "prezzo_unitario": st.column_config.NumberColumn("prezzo_unitario", step=0.01, min_value=0.0),
        },
        disabled=["id","categoria","fornitore","codice_fornitore"],  # non modificabili
        height=600,
        key="mat_editor"
    )

    if st.button("💾 Salva modifiche materiali"):
        # Passo il DF originale delle righe visibili (per confronto)
        orig_for_edited = df_all[df_all["id"].isin(edited["id"])]
        update_materiali_bulk(edited, orig_for_edited)
        st.rerun()
     
    # Anteprima impatti manuale (se ci sono modifiche prezzo recenti)
    if st.session_state.get("last_changed_material_ids"):
        if st.button("🔁 Ricalcola/mostra impatti voci colpite"):
            df_imp = anteprima_impatti_materiali(st.session_state["last_changed_material_ids"])
            st.caption(f"Voci impattate: {len(df_imp)}")
            if df_imp.empty:
                st.info("Nessuna voce legata ai materiali modificati.")
        else:
            st.dataframe(
                df_imp.drop(columns=["voce_id"]),
                use_container_width=True, hide_index=True, height=380
            ) 
            
               # Import CSV/Excel
    # ---------------------------
    with st.expander("📥 Import materiali da CSV/Excel"):
        st.markdown("Colonne richieste: **categoria, fornitore, codice_fornitore, descrizione, unita_misura, prezzo_unitario** (+ opz. `quantita_default`).")
        up = st.file_uploader("Carica file .csv o .xlsx", type=["csv","xlsx"], key="up_materiali")
        if up is not None:
            import_materiali_csv(up)
            st.rerun()
    with st.expander("🕘 Storico prezzi materiali"):
        with get_con() as con:
            df = pd.read_sql_query("""
                SELECT s.changed_at, m.descrizione AS materiale,
                       s.prezzo_vecchio, s.prezzo_nuovo, s.note
                FROM materiali_prezzi_storico s
                JOIN materiali_base m ON m.id = s.materiale_id
                ORDER BY s.changed_at DESC
                LIMIT 200
            """, con)
            st.dataframe(df, use_container_width=True, hide_index=True, height=240)
       

# ------------------------------------------------------------------
# UI – Capitoli
# ------------------------------------------------------------------
def ui_capitoli():
    st.subheader("Capitoli (Categorie lavorazioni)")
    with st.form("form_capitolo"):
        c1, c2 = st.columns([1, 2.2])
        c3, c4 = st.columns(2)
        codice = c1.text_input("Codice capitolo", placeholder="Es. CAP.1")
        nome = c2.text_input("Nome capitolo", placeholder="Es. CANTIERAMENTO")
        cg_def = c3.number_input("Spese generali default (%)", min_value=0.0, value=0.0, step=0.5)
        ut_def = c4.number_input("Utile impresa default (%)", min_value=0.0, value=0.0, step=0.5)
        if st.form_submit_button("➕ Aggiungi capitolo"):
            if not (codice and nome):
                st.warning("Inserisci codice e nome.")
            else:
                add_capitolo(codice, nome, cg_def, ut_def); st.rerun()

    df = df_capitoli()
    st.dataframe(df.rename(columns={"cg_def":"CG% default","ut_def":"Utile% default"}),
                 use_container_width=True, hide_index=True)

    if not df.empty:
        st.divider()
        st.markdown("**Aggiorna default capitolo (influenza nuove voci; le esistenti restano invariate)**")
        cid = st.selectbox("Capitolo", options=[None]+df["id"].tolist(),
                           format_func=lambda x: "—" if x is None else f"{df[df['id']==x]['codice'].iloc[0]} – {df[df['id']==x]['nome'].iloc[0]}")
        if cid:
            row = df[df["id"]==cid].iloc[0]
            c1, c2, c3 = st.columns(3)
            new_cg = c1.number_input("CG% default", min_value=0.0, value=float(row["cg_def"]), step=0.5, key=f"cgdef_{cid}")
            new_ut = c2.number_input("Utile% default", min_value=0.0, value=float(row["ut_def"]), step=0.5, key=f"utdef_{cid}")
            if c3.button("💾 Aggiorna default"):
                update_capitolo_defaults(int(cid), new_cg, new_ut); st.rerun()

        del_id = st.selectbox("Elimina capitolo", options=[None]+df["id"].tolist(),
                              format_func=lambda x: "—" if x is None else f"{df[df['id']==x]['codice'].iloc[0]} – {df[df['id']==x]['nome'].iloc[0]}")
        if del_id and st.button("Elimina"):
            delete_capitolo(int(del_id)); st.rerun()

# ------------------------------------------------------------------
# UI – Voci di analisi (con filtri Capitolo+Descrizione e chiavi coerenti)
# ------------------------------------------------------------------
def ui_voci():
    st.subheader("Voci di analisi")

    cap = df_capitoli()
    mats_all = df_materiali()

    if cap.empty:
        st.info("Crea almeno un capitolo.")
        return
    if mats_all.empty:
        st.info("Aggiungi almeno un materiale nell'Archivio.")
        return

    # -----------------------------
    # FORM: creazione nuova voce
    # -----------------------------
    with st.form("form_voce_new"):
        c1, c2 = st.columns([2, 3])

        cap_map = {int(r.id): f"{r.codice} – {r.nome}" for _, r in cap.iterrows()}
        capitolo_id = c1.selectbox(
            "Capitolo",
            options=list(cap_map.keys()),
            format_func=lambda x: cap_map[x],
            key="capitolo_per_voce"  # key dedicata
        )

        rowc = cap[cap["id"] == capitolo_id].iloc[0]
        cg_def, ut_def = float(rowc["cg_def"]), float(rowc["ut_def"])

        # Elenco voci già presenti nel capitolo selezionato (anti-duplicato)
        voci_cap = df_voci(capitolo_id)
        with st.expander("Voci già presenti in questo capitolo (codice + descrizione)", expanded=False):
            if voci_cap.empty:
                st.caption("Nessuna voce nel capitolo.")
            else:
                st.dataframe(
                    voci_cap[["codice", "descrizione"]].sort_values("codice"),
                    use_container_width=True, hide_index=True, height=220
                )

        codice = c1.text_input("Codice voce", placeholder="Es. 01")
        descrizione = c2.text_input("Descrizione voce", placeholder="Es. Recinzione di cantiere")
        cg_pct = c1.number_input("Spese generali (%)", min_value=0.0, value=cg_def, step=0.5)
        utile_pct = c1.number_input("Utile impresa (%)", min_value=0.0, value=ut_def, step=0.5,
                                    help="Calcolato su (materie + spese generali).")
        um_voce = c2.selectbox("UM della VOCE (misura prodotta)", UM_CHOICES)
        q_voce = c2.number_input("Quantità della VOCE (misura prodotta)", min_value=0.0, value=1.0, step=0.1)

        # Prezzo di riferimento (facoltativo)
        prezzo_rif = c1.number_input("Prezzo di riferimento (facolt.)", min_value=0.0, value=0.0, step=0.01, format="%.2f")

        if st.form_submit_button("➕ Crea voce"):
            if not (codice and descrizione and um_voce and q_voce > 0):
                st.warning("Compila codice, descrizione, UM voce e quantità voce (>0).")
            else:
                add_voce(capitolo_id, codice, descrizione, cg_pct, utile_pct, um_voce, q_voce, prezzo_rif)
                st.rerun()

    # -----------------------------
    # FILTRO per capitolo (dropdown) + filtri stile Excel
    # -----------------------------
    cap_map = {int(r.id): f"{r.codice} – {r.nome}" for _, r in cap.iterrows()}
    filtro_list = st.selectbox(
        "Filtra per capitolo",
        options=[0] + list(cap_map.keys()),
        format_func=lambda x: "Tutti" if x == 0 else cap_map[x],
        key="filtro_cap_list"  # key dedicata
    )

    voci = df_voci(None if filtro_list == 0 else filtro_list)
    if voci.empty:
        st.info("Nessuna voce trovata.")
        return

    # Filtri stile Excel: Capitolo (codice+nome) + Descrizione
    fv1, fv2 = st.columns([1.6, 2])
    f_cap  = fv1.text_input("Filtro Capitolo (codice/nome)", key="flt_voci_cap")
    f_desc = fv2.text_input("Filtro Descrizione", key="flt_voci_desc")

    vv = voci.copy()
    cap_full = vv["capitolo_codice"].astype(str) + " " + vv["capitolo_nome"].astype(str)
    mask = like_mask(cap_full, f_cap) & like_mask(vv["descrizione"], f_desc)
    vv = vv[mask]

    # -----------------------------
    # Lista a sinistra, dettaglio a destra
    # -----------------------------
    left, right = st.columns([2, 3], vertical_alignment="top")

    with left:
        st.write("Voci disponibili")
        st.dataframe(
            vv.rename(columns={
                "cg_pct": "CG %", "utile_pct": "Utile %", "um_voce": "UM Voce", "q_voce": "Q.tà Voce"
            })[
                ["id", "capitolo_codice", "capitolo_nome", "codice", "descrizione", "UM Voce", "Q.tà Voce", "CG %", "Utile %"]
            ],
            use_container_width=True, hide_index=True, height=620  # +scroll (>> 3 righe)
        )

        ids = vv["id"].tolist()
        labels = [f"{r.capitolo_codice} {r.codice} – {r.descrizione[:50]}" for _, r in vv.iterrows()]
        voce_sel = st.selectbox(
            "Seleziona voce",
            options=[None] + ids,
            format_func=lambda x: "—" if x is None else labels[ids.index(x)],
            key="voce_sel_list"  # key dedicata
        )

        colA, colB, _ = st.columns(3)
        if voce_sel and colA.button("🗑️ Elimina voce"):
            delete_voce(int(voce_sel))
            st.rerun()
        if voce_sel and colB.button("🧬 Duplica voce"):
            clone_voce(int(voce_sel))
            st.rerun()

    if voce_sel:
        with right:
            v = get_voce(int(voce_sel))
            st.markdown(f"**Voce:** {v['capitolo_codice']} {v['codice']} – {v['descrizione']}")

            # Edit parametri voce (incl. prezzo riferimento)
            c1, c2, c3, c4, c5 = st.columns(5)
            new_cg = c1.number_input("Spese generali (%)", min_value=0.0, value=float(v["cg_pct"]), step=0.5, key=f"cg_{voce_sel}")
            new_ut = c2.number_input("Utile impresa (%)", min_value=0.0, value=float(v["utile_pct"]), step=0.5, key=f"ut_{voce_sel}")
            new_um = c3.selectbox("UM Voce", UM_CHOICES,
                                  index=UM_CHOICES.index(v["um_voce"]) if v["um_voce"] in UM_CHOICES else 0,
                                  key=f"um_{voce_sel}")
            new_qv = c4.number_input("Q.tà Voce", min_value=0.0, value=float(v["q_voce"]), step=0.1, key=f"qv_{voce_sel}")
            new_prezzo_rif = c5.number_input("Prezzo di riferimento", min_value=0.0,
                                             value=float(v.get("prezzo_rif", 0.0)), step=0.01, format="%.2f")

            if c5.button("💾 Aggiorna voce", key=f"upd_{voce_sel}"):
                update_voce_perc_umqty(int(voce_sel), new_cg, new_ut, new_um, new_qv, new_prezzo_rif)
                st.rerun()

            st.divider()
            st.write("Distinta base – aggiungi riga")

            mats = df_materiali()
            mat_map = {int(r.id): f"{r.categoria} | {r.fornitore} | {r.codice_fornitore} – {str(r.descrizione)[:50]}"
                       for _, r in mats.iterrows()}
            co1, co2, co3 = st.columns([3, 1, 1])

            # piccola ricerca rapida sul select dei materiali
            search = st.text_input("Cerca materiale (categoria/fornitore/codice/descrizione)", key=f"search_{voce_sel}")
            filtered_ids = list(mat_map.keys())
            if search:
                s = search.strip().lower()
                def match(txt): return s in str(txt).lower()
                filtered_ids = [mid for mid in mat_map if match(mat_map[mid])]
                if not filtered_ids:
                    st.info("Nessun materiale corrispondente alla ricerca.")
                    filtered_ids = list(mat_map.keys())

            mat_id = co1.selectbox("Materiale", options=filtered_ids,
                                   format_func=lambda x: mat_map[x], key=f"m_{voce_sel}")
            qta = co2.number_input("Quantità", min_value=0.0, value=1.0, step=0.1, key=f"q_{voce_sel}")
            if co3.button("➕ Aggiungi", key=f"add_{voce_sel}"):
                if qta <= 0:
                    st.warning("Quantità > 0")
                else:
                    add_riga_distinta(int(voce_sel), int(mat_id), float(qta))
                    st.rerun()

            righe = df_righe(int(voce_sel))
            if righe.empty:
                st.info("Nessuna riga in distinta.")
            else:
                view = righe[[
                    "id", "categoria", "fornitore", "codice_fornitore",
                    "materiale_descrizione", "unita_misura", "prezzo_unitario",
                    "quantita", "subtotale"
                ]].copy()
                edited = st.data_editor(
                    view, use_container_width=True, num_rows="fixed",
                    column_config={"quantita": st.column_config.NumberColumn("quantita", step=0.1)}
                )
                colx, coly = st.columns([1, 1])
                if colx.button("💾 Salva quantità modificate", key=f"saveq_{voce_sel}"):
                    update_quantita_righe(int(voce_sel), edited, view)
                    st.rerun()
                rid = coly.selectbox("Elimina riga", options=[None] + righe["id"].tolist(),
                                     format_func=lambda x: "—" if x is None else f"riga #{x}", key=f"delrow_{voce_sel}")
                if rid and st.button("Elimina selezionata", key=f"btn_del_{voce_sel}"):
                    delete_riga(int(rid))
                    st.rerun()

            tot = compute_totali_voce(int(voce_sel))
            m1, m2, m3, m4 = st.columns(4)
            m1.metric("Materie (€)", f"{tot['costo_materie']:.2f}")
            m2.metric("Spese generali (€)", f"{tot['costi_generali']:.2f}", help=f"{tot['cg_pct']}%")
            m3.metric("Utile (€)", f"{tot['utile']:.2f}", help=f"{tot['utile_pct']}%")
            m4.metric("Totale voce (€)", f"{tot['totale']:.2f}")

            # Scostamento % rispetto al prezzo di riferimento (se impostato)
            prezzo_rif_show = float(new_prezzo_rif) if new_prezzo_rif is not None else float(v.get("prezzo_rif", 0.0))
            if prezzo_rif_show > 0:
                delta_pct = (tot["totale"] - prezzo_rif_show) / prezzo_rif_show * 100.0
                st.metric("% scostamento prezzo di riferimento", f"{tot['totale']:.2f} €", delta=f"{delta_pct:+.2f}%")
            else:
                st.caption("Prezzo di riferimento non impostato.")


# ------------------------------------------------------------------
# UI – Sommario EPU (con filtri Capitolo+Descrizione) + export
# ------------------------------------------------------------------
def ui_sommario():
    st.subheader("Sommario EPU")

    voci = df_voci()
    if voci.empty:
        st.info("Nessuna voce disponibile.")
        return

    # -----------------------------
    # Costruzione tabella sintetica
    # -----------------------------
    rows = []
    for _, r in voci.iterrows():
        tot = compute_totali_voce(int(r.id))
        rows.append({
            "Capitolo": r.capitolo_codice,     # codice
            "CapitoloNome": r.capitolo_nome,   # nome
            "Cod. Voce": r.codice,
            "Descrizione": r.descrizione,
            "UM Voce": r.um_voce,
            "Q.tà Voce": r.q_voce,
            "CG %": r.cg_pct,
            "Utile %": r.utile_pct,
            "Materie (€)": round(tot["costo_materie"], 2),
            "Spese generali (€)": round(tot["costi_generali"], 2),
            "Utile (€)": round(tot["utile"], 2),
            "Totale (€)": round(tot["totale"], 2),
            "voce_id": int(r.id),
        })
    df_sum = pd.DataFrame(rows).sort_values(["Capitolo", "Cod. Voce"]).reset_index(drop=True)

    # -----------------------------
    # Filtri "stile Excel"
    # -----------------------------
    f1, f2 = st.columns([1.6, 2])
    filtro_cap = f1.text_input("Filtro Capitolo (codice o nome)", key="flt_som_cap")
    filtro_desc = f2.text_input("Filtro Descrizione", key="flt_som_desc")

    if filtro_cap:
        cap_query = filtro_cap.strip().lower()
        cap_full = (df_sum["Capitolo"].astype(str) + " " + df_sum["CapitoloNome"].astype(str)).str.lower()
        df_sum = df_sum[cap_full.str.contains(cap_query, na=False)]

    if filtro_desc:
        desc_query = filtro_desc.strip().lower()
        df_sum = df_sum[df_sum["Descrizione"].astype(str).str.lower().str.contains(desc_query, na=False)]

    # -----------------------------
    # Tabella principale
    # -----------------------------
    st.caption(f"Voci in elenco: {len(df_sum)}")
    st.dataframe(
        df_sum.drop(columns=["voce_id"]),  # mostro anche Capitolo + Nome
        use_container_width=True,
        hide_index=True,
        height=420  # scrollbar per elenchi lunghi
    )

    # -----------------------------
    # Dettaglio a livelli (con scroll)
    # -----------------------------
    st.markdown("### Dettaglio a livelli")
    for (cap_code, cap_name), grp in df_sum.groupby(["Capitolo", "CapitoloNome"], sort=False):
        with st.expander(f"📁 Capitolo {cap_code} — {cap_name} | Voci: {len(grp)}", expanded=False):
            # elenco voci completo; nessun limite a 3 — verranno mostrate tutte
            for _, r in grp.iterrows():
                titolo_voce = (
                    f"🧩 Voce {r['Cod. Voce']} – {r['Descrizione']} "
                    f"({r['Q.tà Voce']} {r['UM Voce']}) | Totale € {r['Totale (€)']:.2f}"
                )
                with st.expander(titolo_voce, expanded=False):
                    righe = df_righe(int(r["voce_id"]))
                    if righe.empty:
                        st.info("Nessuna riga.")
                    else:
                        show = righe[[
                            "categoria", "fornitore", "codice_fornitore",
                            "materiale_descrizione", "unita_misura",
                            "prezzo_unitario", "quantita", "subtotale"
                        ]]
                        # height alto -> scrollbar interna, utile se tante righe
                        st.dataframe(show, use_container_width=True, hide_index=True, height=320)

                    st.caption(
                        f"Materie € {r['Materie (€)']:.2f} | "
                        f"Spese generali € {r['Spese generali (€)']:.2f} | "
                        f"Utile € {r['Utile (€)']:.2f}"
                    )

    # -----------------------------
    # Export Excel (Sommario EPU)
    # -----------------------------
    buf = export_excel()
    if buf:
        st.download_button(
            "⬇️ Esporta Excel (Sommario EPU)",
            data=buf.getvalue(),
            file_name="EPU_sommario.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

# ------------------------------------------------------------------
# UI – Clienti (riusata nella pagina Preventivi)
# ------------------------------------------------------------------
def ui_clienti():
    st.subheader("Clienti")
    df = df_clienti()
    st.dataframe(df, use_container_width=True, hide_index=True, height=260)

    with st.expander("➕ Nuovo cliente"):
        c1, c2 = st.columns(2)
        c3, c4 = st.columns(2)
        nome = c1.text_input("Nome *")
        piva = c1.text_input("P.IVA / CF")
        indirizzo = c2.text_input("Indirizzo")
        cap_zip = c3.text_input("CAP")  # rinominato da 'cap'
        citta = c3.text_input("Città", key="citta_cli")
        provincia = c4.text_input("Provincia")
        nazione = c4.text_input("Nazione", value="Italia")
        email = c1.text_input("Email", key="email_cli")
        telefono = c2.text_input("Telefono", key="tel_cli")
        note = st.text_area("Note")
        if st.button("Salva cliente") and nome:
            add_cliente(nome=nome, piva=piva, indirizzo=indirizzo, cap=cap_zip, citta=citta, provincia=provincia,
                        nazione=nazione, email=email, telefono=telefono, note=note)
            st.rerun()

    # --- Elimina cliente (solo se senza preventivi) ---
    if not df.empty:
        st.divider()
        cid = st.selectbox("Elimina cliente", options=[None] + df["id"].tolist(),
                           format_func=lambda x: "—" if x is None else df[df["id"]==x]["nome"].iloc[0])
        if cid and st.button("Elimina cliente selezionato"):
            delete_cliente(int(cid))
            st.rerun()

# ------------------------------------------------------------------
# UI – Preventivi
# ------------------------------------------------------------------
def render_preventivo_view(pid: int):
    # Mostra dettagli, totali e export per un preventivo esistente (vista in Archivio)
    testa, righe = df_preventivo(int(pid))
    if testa.empty:
        st.warning("Preventivo non trovato.")
        return

    st.markdown(
        f"#### Preventivo {testa['numero'].iloc[0]} del {testa['data'].iloc[0]} – {testa['cliente_nome'].iloc[0]}"
    )

    if not righe.empty:
        st.dataframe(
            righe[
                [
                    "capitolo_codice",
                    "capitolo_nome",
                    "voce_codice",
                    "descrizione",
                    "um",
                    "quantita",
                    "prezzo_unitario",
                    "prezzo_totale",
                ]
            ],
            use_container_width=True,
            hide_index=True,
        )

        by_cap = (
            righe.groupby(["capitolo_codice", "capitolo_nome"])["prezzo_totale"]
            .sum()
            .reset_index()
            .rename(columns={"prezzo_totale": "Totale (€)"})
        )
        st.markdown("**Totali per capitolo**")
        st.dataframe(by_cap, use_container_width=True, hide_index=True)
    else:
        st.info("Nessuna riga nel preventivo.")

    # Totali documento
    imp = float(testa["imponibile"].iloc[0])
    iva_p = float(testa["iva_percentuale"].iloc[0])
    iva_imp = float(testa["iva_importo"].iloc[0])
    tot = float(testa["totale"].iloc[0])
    c1, c2, c3 = st.columns(3)
    c1.metric("Imponibile (€)", f"{imp:.2f}")
    c2.metric(f"IVA {iva_p:.0f}% (€)", f"{iva_imp:.2f}")
    c3.metric("Totale documento (€)", f"{tot:.2f}")

    # Pulsanti export con KEY univoche (evita StreamlitDuplicateElementId)
    colx, coly = st.columns(2)

    buf_xls = export_preventivo_excel(int(pid))
    if buf_xls:
        colx.download_button(
            "⬇️ Excel",
            data=buf_xls.getvalue(),
            file_name=f"Preventivo_{testa['numero'].iloc[0]}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            key=f"dl_xls_view_{pid}",
        )

    buf_docx = export_preventivo_docx(int(pid))
    if buf_docx:
        coly.download_button(
            "⬇️ Word (DOCX)",
            data=buf_docx.getvalue(),
            file_name=f"Preventivo_{testa['numero'].iloc[0]}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"dl_docx_view_{pid}",
        )

def ui_preventivi():
    st.subheader("Preventivi")

    # 3 tab: Nuovo/Modifica, Clienti, Archivio
    tab1, tab2, tab3 = st.tabs(["🧾 Nuovo/Modifica", "👥 Clienti", "📚 Archivio"])

    # --- Tab Clienti ---
    with tab2:
        # usa la stessa UI clienti che abbiamo già
        ui_clienti()

    # --- Tab Nuovo/Modifica ---
    with tab1:
        cli = df_clienti()
        if cli.empty:
            st.info("Inserisci almeno un Cliente nella tab 'Clienti'.")
            return

        st.markdown("### Testata preventivo")
        c1, c2, c3 = st.columns([2,1,1])
        cliente_id = c1.selectbox("Cliente", options=cli["id"], format_func=lambda i: cli[cli["id"]==i]["nome"].iloc[0])
        numero = c2.text_input("Numero", placeholder="2025-001")
        data = c3.text_input("Data (YYYY-MM-DD)", placeholder="2025-08-12")
        note_finali = st.text_area("Note finali (facoltative)")
        iva_percent = st.number_input("IVA %", min_value=0.0, value=22.0, step=1.0)

        colh1, colh2 = st.columns([1,1])
        if colh1.button("➕ Crea preventivo"):
            if not (numero and data):
                st.warning("Numero e Data sono obbligatori.")
            else:
                pid = create_preventivo(numero, data, int(cliente_id), note_finali, iva_percent)
                st.session_state["preventivo_corrente"] = pid
                st.success(f"Preventivo creato (ID {pid}).")
                st.rerun()

        # Se ho un preventivo corrente, consenti aggiunta righe e totali
        pid = st.session_state.get("preventivo_corrente")
        if pid:
            st.caption(f"Preventivo corrente: ID {pid}")

            st.markdown("### Aggiungi righe")
            cap = df_capitoli()
            if cap.empty:
                st.info("Crea almeno un capitolo e una voce nella sezione Voci di analisi.")
                return
            cap_map = {int(r.id): f"{r.codice} – {r.nome}" for _, r in cap.iterrows()}
            scel_cap = st.selectbox("Capitolo", options=list(cap_map.keys()), format_func=lambda x: cap_map[x])

            voci = df_voci(scel_cap)
            if voci.empty:
                st.info("Nessuna voce disponibile per questo capitolo.")
                return
            voce_map = {int(r.id): f"{r.codice} – {r.descrizione}" for _, r in voci.iterrows()}
            scel_voce = st.selectbox("Voce", options=list(voce_map.keys()), format_func=lambda x: voce_map[x])

            v = get_voce(int(scel_voce))
            prezzo_u = prezzo_unitario_voce(int(scel_voce))
            desc_default = v["descrizione"]
            um_default = v["um_voce"]

            desc = st.text_input("Descrizione riga", value=desc_default)
            note_riga = st.text_area("Note riga (facoltative)")
            um = st.selectbox("UM", UM_CHOICES, index=UM_CHOICES.index(um_default) if um_default in UM_CHOICES else 0)
            qta = st.number_input("Quantità", min_value=0.0, value=1.0, step=0.1)
            st.write(f"**Prezzo unitario (auto)**: € {prezzo_u:.2f}")
            if st.button("➕ Aggiungi riga"):
                if qta <= 0:
                    st.warning("Quantità > 0")
                else:
                    add_riga_preventivo(int(pid), int(scel_cap), int(scel_voce), desc, note_riga, um, qta, prezzo_u)
                    ricalcola_totali_preventivo(int(pid), iva_percent)
                    st.success("Riga aggiunta.")
                    st.rerun()

            # Vista righe + totali
            testa, righe = df_preventivo(int(pid))
            if not righe.empty:
                st.markdown("#### Righe inserite")
                st.dataframe(righe[["capitolo_codice","capitolo_nome","voce_codice","descrizione","um","quantita","prezzo_unitario","prezzo_totale"]],
                             use_container_width=True, hide_index=True)

                st.markdown("#### Totali per capitolo")
                by_cap = (righe.groupby(["capitolo_codice","capitolo_nome"])["prezzo_totale"].sum()
                          .reset_index().rename(columns={"prezzo_totale":"Totale (€)"}))
                st.dataframe(by_cap, use_container_width=True, hide_index=True)

            # Totali documento
            ricalcola_totali_preventivo(int(pid), iva_percent)
            testa, _ = df_preventivo(int(pid))
            imp = float(testa["imponibile"].iloc[0])
            iva_p = float(testa["iva_percentuale"].iloc[0])
            iva_imp = float(testa["iva_importo"].iloc[0])
            tot = float(testa["totale"].iloc[0])

            c1, c2, c3 = st.columns(3)
            c1.metric("Imponibile (€)", f"{imp:.2f}")
            c2.metric(f"IVA {iva_p:.0f}% (€)", f"{iva_imp:.2f}")
            c3.metric("Totale documento (€)", f"{tot:.2f}")

            # --- Azioni finali ---
            if c3.button("💾 Salva e archivia"):
                ricalcola_totali_preventivo(int(pid), iva_percent)
                st.session_state["last_saved_preventivo_id"] = int(pid)
                st.session_state.pop("preventivo_corrente", None)
                st.success(f"Preventivo {testa['numero'].iloc[0]} salvato e archiviato. Vai nella tab 'Archivio' per vederlo.")
                st.rerun()

    # --- Tab Archivio ---
    with tab3:
        st.markdown("### Archivio preventivi")
        cli = df_clienti()

        colf1, colf2, colf3, colf4 = st.columns([1,1,1,1])
        numero_like = colf1.text_input("Filtra per numero")
        data_like = colf2.text_input("Filtra per data (YYYY-MM-DD)")
        cli_sel = colf3.selectbox("Cliente", options=[0]+cli["id"].tolist(),
                                  format_func=lambda x: "Tutti" if x==0 else cli[cli["id"]==x]["nome"].iloc[0])
        if colf4.button("🔄 Aggiorna elenco"):
            st.rerun()

        arch = df_preventivi_archivio(numero_like, data_like, None if cli_sel==0 else cli_sel)
        st.caption(f"{len(arch)} preventivi trovati")
        st.dataframe(arch, use_container_width=True, hide_index=True)

        # Evidenzia l’ultimo salvato (se presente in sessione)
        last_id = st.session_state.get("last_saved_preventivo_id")
        if last_id is not None and not arch.empty and (arch["id"] == last_id).any():
            st.success(f"Ultimo salvato: ID {last_id}")
            default_idx = [None] + arch["id"].tolist()
            preselect = default_idx.index(last_id) if last_id in arch["id"].tolist() else 0
        else:
            default_idx = [None] + arch["id"].tolist()
            preselect = 0

        # Apri preventivo
        opened_id = st.session_state.get("opened_preventivo_from_archivio")
        pid_open = st.selectbox(
            "Apri preventivo",
            options=[None] + arch["id"].tolist(),
            index=preselect if opened_id is None else ([None] + arch["id"].tolist()).index(opened_id) if opened_id in arch["id"].tolist() else 0,
            format_func=lambda x: "—" if x is None else f"ID {x}"
        )

        if pid_open and st.button("Apri"):
            st.session_state["preventivo_corrente"] = int(pid_open)  # utile anche per la tab Nuovo/Modifica
            st.session_state["opened_preventivo_from_archivio"] = int(pid_open)
            st.success(f"Preventivo ID {int(pid_open)} aperto qui sotto.")
            st.rerun()

        # Se c'è un preventivo selezionato/precedente, mostrane i dettagli QUI
        pid_show = st.session_state.get("opened_preventivo_from_archivio")
        if pid_show:
            st.divider()
            render_preventivo_view(int(pid_show))

# ------------------------------------------------------------------
# MAIN
# ------------------------------------------------------------------
def main():
    init_db()
    st.title("🏗️ EPU Builder v1.3.2")
    st.caption("Analisi voci (CG%/Utile% capitolo), distinte, Sommario EPU, preventivi con export Excel/Word.")

    pagina = st.sidebar.radio("Navigazione", [
        "Categorie", "Fornitori", "Archivio materiali", "Capitoli", "Voci di analisi", "Sommario EPU", "Preventivi"
    ])

    if pagina == "Categorie":
        ui_categorie()
    elif pagina == "Fornitori":
        ui_fornitori()
    elif pagina == "Archivio materiali":
        ui_materiali()
    elif pagina == "Capitoli":
        ui_capitoli()
    elif pagina == "Voci di analisi":
        ui_voci()
    elif pagina == "Sommario EPU":
        ui_sommario()
    elif pagina == "Preventivi":
        ui_preventivi()

if __name__ == "__main__":
    main()
