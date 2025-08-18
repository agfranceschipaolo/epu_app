# autopatch.py — applica patch (2)(3)(4) a App.py in modo idempotente
# Crea un backup App.py.bak.<timestamp>
import re, shutil, time
from pathlib import Path

FILE = Path("App.py")
orig = FILE.read_text(encoding="utf-8")

def backup():
    ts = time.strftime("%Y%m%d-%H%M%S")
    bak = FILE.with_suffix(f".py.bak.{ts}")
    shutil.copy2(FILE, bak)
    print(f"🗄️  Backup creato: {bak.name}")

def ensure(pattern, text):
    return re.search(pattern, text, flags=re.S) is not None

def insert_after(anchor_regex, snippet, text, tag=""):
    m = re.search(anchor_regex, text, flags=re.S)
    if not m:
        raise RuntimeError(f"Anchor non trovato per: {tag or anchor_regex}")
    pos = m.end()
    return text[:pos] + snippet + text[pos:]

def replace_block(start_regex, end_regex, new_block, text, tag=""):
    m1 = re.search(start_regex, text, flags=re.S)
    if not m1:
        raise RuntimeError(f"Inizio blocco non trovato: {tag or start_regex}")
    m2 = re.search(end_regex, text[m1.start():], flags=re.S)
    if not m2:
        raise RuntimeError(f"Fine blocco non trovato: {tag or end_regex}")
    s = m1.start()
    e = m1.start() + m2.end()
    return text[:s] + new_block + text[e:]

def add_storico_table(text):
    if ensure(r"CREATE TABLE IF NOT EXISTS materiali_prezzi_storico", text):
        return text, False
    # inserisci tabella storico subito dopo preventivo_righe
    anchor = r'CREATE TABLE IF NOT EXISTS preventivo_righe .*?\)\s*"""\)\)'
    snippet = r"""

        # Storico prezzi materiali
        cur.execute("""
        CREATE TABLE IF NOT EXISTS materiali_prezzi_storico (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            materiale_id INTEGER NOT NULL,
            prezzo_vecchio REAL NOT NULL,
            prezzo_nuovo REAL NOT NULL,
            changed_at TEXT NOT NULL DEFAULT (datetime('now')),
            note TEXT,
            FOREIGN KEY(materiale_id) REFERENCES materiali_base(id)
        )""")
"""
    new_text = insert_after(anchor, snippet, text, tag="storico-table")
    # indici
    if not ensure(r"idx_sto_mat", new_text):
        idx_anchor = r'cur\.execute\("CREATE INDEX IF NOT EXISTS idx_prev_data ON preventivi\(data\)"\)'
        idx_snip = '\n        cur.execute("CREATE INDEX IF NOT EXISTS idx_sto_mat ON materiali_prezzi_storico(materiale_id)")\n        cur.execute("CREATE INDEX IF NOT EXISTS idx_sto_date ON materiali_prezzi_storico(changed_at)")'
        new_text = insert_after(idx_anchor, idx_snip, new_text, tag="storico-index")
    return new_text, True

def add_impatti_functions(text):
    if ensure(r"def anteprima_impatti_materiali", text):
        return text, False
    anchor = r"def prezzo_unitario_voce\(voce_id: int\).*?return tot / q"
    snippet = r"""

# -------- (2) Impatti da aggiornamento materiali --------
def voci_impattate_da_materiali(material_ids: list[int]) -> pd.DataFrame:
    """ + '"""' + r"""Voci che usano almeno uno dei materiali indicati.""" + '"""' + r"""
    if not material_ids:
        return pd.DataFrame(columns=["voce_id","capitolo_codice","capitolo_nome","codice","descrizione"])
    with get_con() as con:
        q = """
        SELECT DISTINCT v.id AS voce_id, c.codice AS capitolo_codice, c.nome AS capitolo_nome,
                        v.codice, v.descrizione, IFNULL(v.prezzo_riferimento,0.0) AS prezzo_rif
        FROM righe_distinta r
        JOIN voci_analisi v ON v.id = r.voce_analisi_id
        JOIN capitoli c ON c.id = v.capitolo_id
        WHERE r.materiale_id IN ({})
        ORDER BY c.codice, v.codice
        """.format(",".join(["?"]*len(material_ids)))
        return pd.read_sql_query(q, con, params=material_ids)

def anteprima_impatti_materiali(material_ids: list[int]) -> pd.DataFrame:
    """ + '"""' + r"""Totale voce ricalcolato + scostamento vs prezzo_riferimento.""" + '"""' + r"""
    voci_df = voci_impattate_da_materiali(material_ids)
    rows = []
    for _, r in voci_df.iterrows():
        tot = compute_totali_voce(int(r.voce_id))["totale"]
        rif = float(r.get("prezzo_rif", 0.0))
        delta_pct = ((tot - rif) / rif * 100.0) if rif > 0 else None
        rows.append({
            "Capitolo": r.capitolo_codice,
            "Voce": r.codice,
            "Descrizione": r.descrizione,
            "Totale attuale (€)": round(tot, 2),
            "Prezzo riferimento (€)": (round(rif, 2) if rif > 0 else "-"),
            "Δ vs riferimento (%)": (f"{delta_pct:+.2f}%" if delta_pct is not None else "-"),
            "voce_id": int(r.voce_id),
        })
    return pd.DataFrame(rows)
"""
    return insert_after(anchor, snippet, text, tag="impatti-funcs"), True

def patch_update_materiali_bulk(text):
    if ensure(r"materiali_prezzi_storico", text) and ensure(r"price_changes", text):
        return text, False
    # sostituisci l'intera funzione update_materiali_bulk
    start = r"def update_materiali_bulk\(df_edit: pd\.DataFrame, df_orig: pd\.DataFrame\):"
    end = r"\n\ndef "
    body = r"""
def update_materiali_bulk(df_edit: pd.DataFrame, df_orig: pd.DataFrame):
    changes = []
    price_changes = []  # (materiale_id, old_price, new_price)
    for _, row in df_edit.iterrows():
        orig = df_orig[df_orig["id"] == row["id"]].iloc[0]
        fields = ["descrizione", "unita_misura", "quantita_default", "prezzo_unitario"]
        updates = {f: row[f] for f in fields if str(row[f]) != str(orig[f])}
        if updates:
            changes.append((int(row["id"]), updates))
            if "prezzo_unitario" in updates:
                try:
                    old_p = float(orig["prezzo_unitario"])
                    new_p = float(updates["prezzo_unitario"])
                    if old_p != new_p:
                        price_changes.append((int(row["id"]), old_p, new_p))
                except Exception:
                    pass
    if not changes:
        st.info("Nessuna modifica da salvare.")
        return
    with get_con() as con:
        for mid, upd in changes:
            sets = ", ".join([f"{k}=?" for k in upd.keys()])
            vals = list(upd.values()) + [mid]
            _exec(con, f"UPDATE materiali_base SET {sets} WHERE id=?", vals)
        for mid, old_p, new_p in price_changes:
            _exec(con, """INSERT INTO materiali_prezzi_storico (materiale_id, prezzo_vecchio, prezzo_nuovo, note)
                          VALUES (?,?,?,?)""", (mid, float(old_p), float(new_p), "Aggiornamento da editor materiali"))
        con.commit()
    if price_changes:
        st.session_state["last_changed_material_ids"] = [mid for (mid, _, _) in price_changes]
        st.success(f"Salvate {len(changes)} modifiche (di cui {len(price_changes)} su prezzo).")
        with st.expander("Anteprima impatti (voci toccate dai materiali aggiornati)"):
            df_imp = anteprima_impatti_materiali(st.session_state["last_changed_material_ids"])
            st.caption(f"Voci impattate: {len(df_imp)}")
            if not df_imp.empty:
                st.dataframe(df_imp.drop(columns=["voce_id"]), use_container_width=True, hide_index=True, height=320)
    else:
        st.success(f"Salvate {len(changes)} modifiche.")
"""
    try:
        return replace_block(start, end, body + "\n\n", text, tag="update_materiali_bulk"), True
    except RuntimeError:
        # fallback: se non trova il blocco, non falliamo
        return text, False

def patch_ui_materiali(text):
    # inserisci i due blocchi UI se mancano
    if not ensure(r"🔁 Ricalcola/mostra impatti voci colpite", text):
        anchor = r"update_materiali_bulk\(edited, orig_for_edited\)\s*\n\s*st\.rerun\(\)"
        snippet = r"""

    # Anteprima impatti manuale (se ci sono modifiche prezzo recenti)
    if st.session_state.get("last_changed_material_ids"):
        if st.button("🔁 Ricalcola/mostra impatti voci colpite"):
            df_imp = anteprima_impatti_materiali(st.session_state["last_changed_material_ids"])
            st.caption(f"Voci impattate: {len(df_imp)}")
            if df_imp.empty:
                st.info("Nessuna voce legata ai materiali modificati.")
            else:
                st.dataframe(df_imp.drop(columns=["voce_id"]), use_container_width=True, hide_index=True, height=380)
"""
        text = insert_after(anchor, snippet, text, tag="ui-materiali-impatti")
    if not ensure(r"🕘 Storico prezzi materiali", text):
        anchor2 = r"import_materiali_csv\(up\)\s*\n\s*st\.rerun\(\)"
        snippet2 = r"""

    with st.expander("🕘 Storico prezzi materiali"):
        with get_con() as con:
            sto = pd.read_sql_query("""
                SELECT s.changed_at, m.id AS materiale_id, m.descrizione, m.codice_fornitore,
                       s.prezzo_vecchio, s.prezzo_nuovo, s.note
                FROM materiali_prezzi_storico s
                JOIN materiali_base m ON m.id = s.materiale_id
                ORDER BY s.changed_at DESC
            """, con)
        st.dataframe(sto, use_container_width=True, hide_index=True, height=280)
"""
        text = insert_after(anchor2, snippet2, text, tag="ui-materiali-storico")
    return text, True

def replace_docx_function(text):
    if ensure(r"def export_preventivo_docx\(pid: int, logo_path", text):
        return text, False
    start = r"def export_preventivo_docx\(pid: int\):"
    end = r"\n\ndef "
    body = r"""def export_preventivo_docx(pid: int, logo_path: typing.Optional[str] = None):
    from docx import Document
    from docx.shared import Pt, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    testa, righe = df_preventivo(pid)
    if testa.empty:
        return None
    d = Document()
    if logo_path:
        try:
            hdr = d.sections[0].header
            p = hdr.paragraphs[0]
            run = p.add_run()
            run.add_picture(logo_path, width=Inches(1.2))
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        except Exception:
            pass
    title = d.add_paragraph()
    run = title.add_run(f"Preventivo {testa['numero'].iloc[0]} del {testa['data'].iloc[0]}")
    run.bold = True
    table = d.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    headers = ["Capitolo", "Voce", "Descrizione", "UM", "Q.tà", "Prezzo U (€)", "Totale (€)"]
    for i, h in enumerate(headers):
        hdr[i].text = h
    for cell in table.rows[0].cells:
        for paragraph in cell.paragraphs:
            if paragraph.runs:
                paragraph.runs[0].font.bold = True
    for _, r in righe.iterrows():
        row = table.add_row().cells
        row[0].text = f"{r['capitolo_codice']} {r['capitolo_nome']}"
        row[1].text = str(r["voce_codice"])
        row[2].text = str(r["descrizione"])
        row[3].text = str(r["um"])
        row[4].text = f"{r['quantita']:.2f}"
        row[5].text = f"{r['prezzo_unitario']:.2f}"
        row[6].text = f"{r['prezzo_totale']:.2f}"
        note_val = r["note"] if "note" in r and pd.notna(r["note"]) and str(r["note"]).strip() else ""
        if note_val:
            p = d.add_paragraph(f"Note: {note_val}")
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(2)
    if not righe.empty:
        bycap = (righe.groupby(["capitolo_codice","capitolo_nome"])["prezzo_totale"].sum()
                 .reset_index().rename(columns={"prezzo_totale":"Totale capitolo (€)"}))
        d.add_paragraph("Totali per capitolo:")
        for _, rr in bycap.iterrows():
            p = d.add_paragraph(f"- {rr['capitolo_codice']} {rr['capitolo_nome']}: € {rr['Totale capitolo (€)']:.2f}")
            p.paragraph_format.space_after = Pt(2)
    imp = float(testa["imponibile"].iloc[0]); iva_p = float(testa["iva_percentuale"].iloc[0]); iva_imp = float(testa["iva_importo"].iloc[0]); tot = float(testa["totale"].iloc[0])
    for label, val in [("Imponibile", f"€ {imp:.2f}"), (f"IVA {iva_p:.0f}%", f"€ {iva_imp:.2f}"), ("Totale documento", f"€ {tot:.2f}")]:
        p = d.add_paragraph(); run_label = p.add_run(f"{label}: "); run_label.bold = True; p.add_run(val)
    import io
    buf = io.BytesIO(); d.save(buf); buf.seek(0); return buf
"""
    # garantiamo import typing se non c'è
    if not ensure(r"\bimport typing\b", text):
        text = "import typing\n" + text
    try:
        return replace_block(start, end, body + "\n\n", text, tag="docx-func"), True
    except RuntimeError:
        return text, False

def patch_render_preventivo_view(text):
    if ensure(r'file_uploader\("Logo \(opz\.\)"', text):
        return text, False
    anchor = r"def render_preventivo_view\(pid: int\):.*?colx, coly = st\.columns\(2\)"
    # sostituisce blocco export DOCX dentro la view
    repl_pat = r"""buf_docx = export_preventivo_docx\(int\(pid\)\)[\s\S]*?download_button\([\s\S]*?dl_docx_view_"""
    new_block = r"""
    logo_file = coly.file_uploader("Logo (opz.)", type=["png","jpg","jpeg"], key=f"logo_{pid}")
    buf_docx = None
    if coly.button("📄 Genera DOCX", key=f"mk_docx_{pid}"):
        logo_path = None
        if logo_file is not None:
            tmp_path = f"/tmp/logo_{pid}_{logo_file.name}"
            with open(tmp_path, "wb") as f:
                f.write(logo_file.getbuffer())
            logo_path = tmp_path
        buf_docx = export_preventivo_docx(int(pid), logo_path=logo_path)
        if buf_docx:
            st.session_state[f"docxbuf_{pid}"] = buf_docx.getvalue()
            st.success("DOCX generato.")
    if st.session_state.get(f"docxbuf_{pid}"):
        coly.download_button(
            "⬇️ Scarica DOCX",
            data=st.session_state[f"docxbuf_{pid}"],
            file_name=f"Preventivo_{testa['numero'].iloc[0]}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key=f"dl_docx_view_{pid}",
        )
"""
    m = re.search(repl_pat, text)
    if not m:
        # prova sostituzione manuale ancorata al coly.download_button precedente
        text = re.sub(r"buf_docx = export_preventivo_docx[\s\S]+?dl_docx_view_\{pid\}\"\),\s*\)\s*", new_block, text)
        if not ensure(r'file_uploader\("Logo \(opz\.\)"', text):
            raise RuntimeError("Non sono riuscito a sostituire il blocco DOCX in render_preventivo_view")
        return text, True
    s, e = m.span()
    return text[:s] + new_block, True

def add_crud_preventivi(text):
    if ensure(r"def delete_preventivo\(pid: int\):", text):
        return text, False
    anchor = r"def delete_materiale\(mid: int\):[\s\S]*?st\.success\(\"Materiale eliminato\."\)\s*"
    snippet = r"""

# -------- (4) CRUD aggiuntivo preventivi/righe --------
def delete_preventivo(pid: int):
    with get_con() as con:
        _exec(con, "DELETE FROM preventivo_righe WHERE preventivo_id=?", (pid,))
        _exec(con, "DELETE FROM preventivi WHERE id=?", (pid,))
        con.commit()
    st.success(f"Preventivo ID {pid} eliminato.")

def update_preventivo_header(pid: int, numero: str, data_iso: str, cliente_id: int, note_finali: str, iva_percent: float):
    with get_con() as con:
        _exec(con, """UPDATE preventivi
                      SET numero=?, data=?, cliente_id=?, note_finali=?, iva_percentuale=?
                      WHERE id=?""",
              (numero.strip(), data_iso, int(cliente_id), note_finali, float(iva_percent), int(pid)))
        con.commit()
    ricalcola_totali_preventivo(int(pid), iva_percent)
    st.success("Testata preventivo aggiornata.")

def update_riga_preventivo(riga_id: int, descrizione: str, note: str, um: str, quantita: float, prezzo_unitario: float):
    tot = float(quantita) * float(prezzo_unitario)
    with get_con() as con:
        _exec(con, """UPDATE preventivo_righe
                      SET descrizione=?, note=?, um=?, quantita=?, prezzo_unitario=?, prezzo_totale=?
                      WHERE id=?""",
              (descrizione.strip(), note, um, float(quantita), float(prezzo_unitario), tot, int(riga_id)))
        con.commit()
    st.success("Riga preventivo aggiornata.")

def delete_riga_preventivo(riga_id: int):
    with get_con() as con:
        _exec(con, "DELETE FROM preventivo_righe WHERE id=?", (int(riga_id),))
        con.commit()
    st.success("Riga preventivo eliminata.")
"""
    return insert_after(anchor, snippet, text, tag="crud-prev"), True

def patch_ui_preventivi(text):
    changed = False
    # (A) selettore rapido apri esistente
    if not ensure(r"Oppure apri un preventivo esistente", text):
        anchor = r"iva_percent = st\.number_input\(.*?\)\s*"
        snippet = r"""
        # Carica preventivo esistente
        arch_quick = df_preventivi_archivio()
        _colL, _colR = st.columns([2,1])
        pid_pick = _colL.selectbox("Oppure apri un preventivo esistente", options=[0]+arch_quick["id"].tolist(),
                                   format_func=lambda x: "—" if x==0 else f"ID {x}")
        if _colR.button("Apri selezionato") and pid_pick:
            st.session_state["preventivo_corrente"] = int(pid_pick)
            st.success(f"Aperto preventivo ID {int(pid_pick)} per modifica.")
            st.rerun()
"""
        text = insert_after(anchor, snippet, text, tag="prev-quick-open")
        changed = True
    # (B) header edit + delete
    if not ensure(r"💾 Aggiorna testata", text):
        anchor2 = r"st\.caption\(f\"Preventivo corrente: ID \{pid\}\"\)"
        snippet2 = r"""
            # Precompila header con i valori correnti
            testa_cur, _righe_cur = df_preventivo(int(pid))
            if not testa_cur.empty:
                numero = st.text_input("Numero", value=testa_cur["numero"].iloc[0], key=f"num_{pid}")
                data = st.text_input("Data (YYYY-MM-DD)", value=testa_cur["data"].iloc[0], key=f"date_{pid}")
                cliente_id = st.selectbox("Cliente", options=cli["id"],
                                          index=cli.index[cli["id"]==testa_cur["cliente_id"].iloc[0]][0],
                                          format_func=lambda i: cli[cli["id"]==i]["nome"].iloc[0], key=f"cli_{pid}")
                note_finali = st.text_area("Note finali (facoltative)", value=testa_cur["note_finali"].iloc[0] or "", key=f"nf_{pid}")
                iva_percent = st.number_input("IVA %", min_value=0.0, value=float(testa_cur["iva_percentuale"].iloc[0]), step=1.0, key=f"iva_{pid}")

                c_upd1, c_upd2, c_upd3 = st.columns([1,1,1])
                if c_upd1.button("💾 Aggiorna testata"):
                    update_preventivo_header(int(pid), numero, data, int(cliente_id), note_finali, float(iva_percent))
                    st.rerun()
                if c_upd2.button("🗑️ Elimina preventivo"):
                    delete_preventivo(int(pid))
                    st.session_state.pop("preventivo_corrente", None)
                    st.rerun()
"""
        text = insert_after(anchor2, snippet2, text, tag="prev-header-edit")
        changed = True
    # (C) editor righe
    if not ensure(r"#### Modifica riga esistente", text):
        anchor3 = r"st\.dataframe\(righe\[\[\"capitolo_codice\".*?\]\],\s*use_container_width=True,\s*hide_index=True\)\s*"
        snippet3 = r"""

                # Editor righe semplice (seleziona → modifica/elimina)
                rid_opts = righe["id"].tolist()
                if rid_opts:
                    st.markdown("#### Modifica riga esistente")
                    colr1, colr2 = st.columns([2,1])
                    rid_sel = colr1.selectbox("Riga", options=rid_opts, format_func=lambda x: f"id {x}")
                    if rid_sel:
                        rsel = righe[righe["id"]==rid_sel].iloc[0]
                        desc_e = st.text_input("Descrizione", value=str(rsel["descrizione"]), key=f"desc_e_{rid_sel}")
                        note_e = st.text_area("Note (facolt.)", value=str(rsel["note"] or ""), key=f"note_e_{rid_sel}")
                        um_e = st.selectbox("UM", UM_CHOICES,
                                            index=UM_CHOICES.index(rsel["um"]) if rsel["um"] in UM_CHOICES else 0,
                                            key=f"um_e_{rid_sel}")
                        q_e = st.number_input("Quantità", min_value=0.0, value=float(rsel["quantita"]), step=0.1, key=f"q_e_{rid_sel}")
                        pu_e = st.number_input("Prezzo unitario (€)", min_value=0.0, value=float(rsel["prezzo_unitario"]), step=0.01, format="%.2f", key=f"pu_e_{rid_sel}")
                        colb1, colb2 = st.columns([1,1])
                        if colb1.button("💾 Aggiorna riga", key=f"upd_r_{rid_sel}"):
                            update_riga_preventivo(int(rid_sel), desc_e, note_e, um_e, q_e, pu_e)
                            ricalcola_totali_preventivo(int(pid), iva_percent)
                            st.rerun()
                        if colb2.button("🗑️ Elimina riga", key=f"del_r_{rid_sel}"):
                            delete_riga_preventivo(int(rid_sel))
                            ricalcola_totali_preventivo(int(pid), iva_percent)
                            st.rerun()
"""
        text = insert_after(anchor3, snippet3, text, tag="prev-rows-edit")
        changed = True
    # (D) archivio: elimina / apri per modifica
    if not ensure(r"Elimina preventivo selezionato", text):
        anchor4 = r"st\.success\(f\"Preventivo ID \{int\(pid_open\)\} aperto qui sotto\.\"\)\s*st\.rerun\(\)"
        snippet4 = r"""

        # Elimina o apri per modifica
        if pid_open:
            cdel1, cdel2 = st.columns([1,1])
            if cdel1.button("🗑️ Elimina preventivo selezionato"):
                delete_preventivo(int(pid_open))
                st.rerun()
            if cdel2.button("✏️ Apri per modifica in 'Nuovo/Modifica'"):
                st.session_state["preventivo_corrente"] = int(pid_open)
                st.success("Aperto in 'Nuovo/Modifica'. Vai alla tab corrispondente.")
"""
        text = insert_after(anchor4, snippet4, text, tag="prev-arch-actions")
        changed = True
    return text, changed

# ---------------- run ----------------
backup()
text = orig

changed_any = False

for step in (
    add_storico_table,
    add_impatti_functions,
    patch_update_materiali_bulk,
    patch_ui_materiali,
    replace_docx_function,
    patch_render_preventivo_view,
    add_crud_preventivi,
    patch_ui_preventivi,
):
    text, changed = step(text)
    changed_any = changed_any or changed
    print(f"✔ {step.__name__}: {'OK' if changed else 'skip (già presente)'}")

if changed_any:
    FILE.write_text(text, encoding="utf-8")
    print("✅ Patch applicata a App.py")
else:
    print("ℹ️ Nessuna modifica necessaria (già patchato).")
